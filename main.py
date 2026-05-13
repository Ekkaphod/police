import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.shared import RGBColor
from docx.oxml.ns import qn as _qn
from docx.oxml import OxmlElement as _OxmlElem
import datetime
from io import BytesIO
import pandas as pd
import base64
import os
import json
import plotly.express as px
import plotly.graph_objects as go
import numpy as np
import pytz
import copy
import calendar
from collections import defaultdict   # ← เพิ่ม
import traceback as _tb

# =================================================================
# PAGE CONFIG
# =================================================================
st.set_page_config(
    page_title="ระบบบันทึกจับกุม",
    page_icon="Thai_Police.png",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =================================================================
# USERS
# =================================================================
USERS = {
    "admin":      {"password": "123",  "fullname": "Admin"},
    "somkid":     {"password": "1234", "fullname": "ร.ต.อ.สมคิด เชื้อเวียง"},
    "chaiporn":   {"password": "1234", "fullname": "ร.ต.อ.ชัยพร ชอบงาม"},
    "teerawat":   {"password": "1234", "fullname": "ส.ต.อ.ธีระวัฒน์ แก่นสาร"},
    "narongrit":  {"password": "1234", "fullname": "ส.ต.อ.ณรงค์ฤทธิ์ เหล่าดี"},
    "wichakorn":  {"password": "1234", "fullname": "ส.ต.อ.วิชชากร วงษ์โท"},
    "watchara":   {"password": "1234", "fullname": "จ.ส.ต.วัชระ จันสุตะ"},
    "wachirawit": {"password": "1234", "fullname": "ส.ต.ท.วชิรวิชญ์ นันทรักษ์"},
    "wisarut":    {"password": "1234", "fullname": "ส.ต.ท.วิศรุต จันทร์สิงห์"},
    "adisorn":    {"password": "1234", "fullname": "ส.ต.ต.อดิศร ศุภนิกร"},
}

# =================================================================
# SESSION STATE INIT
# =================================================================
for key, default in [
    ("settings", {"theme": "light", "logo": None}),
    ("form", {}),
    ("tab", 0),
    ("records", []),
    ("doc_records", []),
    ("doc_running_numbers", {}),
    ("schedule_data", {}),
]:
    if key not in st.session_state:
        st.session_state[key] = default

# =================================================================
# CSS
# =================================================================
KTB_BLUE       = "#00AEEF"
KTB_BLUE_DARK  = "#0076B6"
KTB_NAVY       = "#003B6F"
KTB_GOLD       = "#F5A623"
KTB_TEXT       = "#1A2B45"
KTB_SURFACE    = "#F5F8FB"
KTB_BORDER     = "#D6E4F0"

st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Sarabun:wght@300;400;500;600;700;800&display=swap');
:root {{
    --ktb-blue:{KTB_BLUE};--ktb-blue-dark:{KTB_BLUE_DARK};--ktb-navy:{KTB_NAVY};
    --ktb-gold:{KTB_GOLD};--ktb-text:{KTB_TEXT};--ktb-surface:{KTB_SURFACE};
    --ktb-border:{KTB_BORDER};--ktb-radius:12px;--ktb-radius-lg:18px;
    --ktb-shadow:0 2px 12px rgba(0,174,239,0.08);--ktb-shadow-md:0 4px 24px rgba(0,59,111,0.12);
    --bg-page:#EBF4FB;--bg-card:#FFFFFF;--text-primary:{KTB_TEXT};
    --text-muted:#5A7A99;--text-label:#2E4D6B;--border-color:{KTB_BORDER};
    --input-bg:#FFFFFF;--sidebar-bg:{KTB_NAVY};
}}
html,body,[class*="css"],.stApp{{font-family:'Sarabun',sans-serif!important;background:var(--bg-page)!important;color:var(--text-primary)!important;}}
header[data-testid="stHeader"],[data-testid="stToolbar"],[data-testid="stDeployButton"]{{display:none!important;}}
.block-container{{max-width:1400px!important;padding:1.5rem 1.5rem 3rem!important;}}
[data-testid="stSidebar"]{{background:var(--sidebar-bg)!important;border-right:1px solid rgba(0,174,239,0.15)!important;min-width:270px!important;max-width:270px!important;}}
[data-testid="stSidebar"] *{{color:#E8F4FD!important;font-family:'Sarabun',sans-serif!important;}}
[data-testid="stSidebar"] .stButton>button{{background:rgba(0,174,239,0.12)!important;border:1px solid rgba(0,174,239,0.3)!important;color:#E8F4FD!important;border-radius:var(--ktb-radius)!important;font-size:14px!important;padding:10px 14px!important;text-align:left!important;transition:background 0.2s,border-color 0.2s!important;}}
[data-testid="stSidebar"] .stButton>button:hover{{background:rgba(0,174,239,0.25)!important;border-color:var(--ktb-blue)!important;}}
div.stButton>button{{width:100%!important;background:linear-gradient(135deg,var(--ktb-blue) 0%,var(--ktb-blue-dark) 100%)!important;color:#fff!important;border:none!important;border-radius:var(--ktb-radius)!important;padding:12px 20px!important;font-size:15px!important;font-weight:600!important;font-family:'Sarabun',sans-serif!important;letter-spacing:0.2px!important;transition:opacity 0.2s,transform 0.15s!important;box-shadow:0 3px 10px rgba(0,118,182,0.25)!important;}}
div.stButton>button:hover{{opacity:0.92!important;transform:translateY(-1px)!important;}}
.stTextInput input,.stDateInput input,.stTimeInput input,.stTextArea textarea,.stNumberInput input,.stSelectbox div[data-baseweb="select"]>div{{background:var(--input-bg)!important;color:var(--text-primary)!important;border:1.5px solid var(--border-color)!important;border-radius:10px!important;font-family:'Sarabun',sans-serif!important;font-size:15px!important;}}
.ktb-card{{background:var(--bg-card);border:1px solid var(--border-color);border-radius:var(--ktb-radius-lg);padding:1.25rem 1.5rem;box-shadow:var(--ktb-shadow);margin-bottom:1rem;}}
.kpi-card{{background:var(--bg-card);border:1px solid var(--border-color);border-left:4px solid var(--ktb-blue);border-radius:var(--ktb-radius);padding:1.1rem 1.25rem;box-shadow:var(--ktb-shadow);}}
.kpi-number{{font-size:36px;font-weight:800;color:var(--ktb-blue)!important;line-height:1.1;margin:4px 0;}}
.kpi-label{{font-size:13px;color:var(--text-muted)!important;font-weight:500;text-transform:uppercase;letter-spacing:0.5px;}}
.kpi-icon{{font-size:22px;margin-bottom:4px;}}
.main-header{{background:linear-gradient(135deg,var(--ktb-blue) 0%,var(--ktb-blue-dark) 60%,var(--ktb-navy) 100%);padding:18px 24px;border-radius:var(--ktb-radius-lg);color:#fff!important;font-size:20px;font-weight:700;display:flex;align-items:center;gap:12px;box-shadow:0 4px 20px rgba(0,59,111,0.2);margin-bottom:1.25rem;}}
.main-header *{{color:#fff!important;}}
.sidebar-profile{{background:rgba(0,174,239,0.1);border:1px solid rgba(0,174,239,0.2);border-radius:var(--ktb-radius);padding:12px 14px;margin:0 0 1rem;font-size:13px;line-height:1.7;}}
.metric-card{{background:var(--bg-card);border:1px solid var(--border-color);border-radius:var(--ktb-radius);padding:1rem 1.1rem;text-align:center;}}
.metric-icon{{font-size:24px;margin-bottom:6px;}}
.metric-no{{font-size:32px;font-weight:800;color:var(--ktb-blue)!important;}}
.metric-text{{font-size:13px;color:var(--text-muted)!important;font-weight:500;}}
</style>
""", unsafe_allow_html=True)


# =================================================================
# DOCX TEMPLATE HELPERS
# =================================================================
FONT_NAME       = "TH SarabunPSK"
FONT_SIZE       = Pt(14)
FONT_SIZE_TITLE = Pt(16)


def _set_font(run, size=FONT_SIZE, bold=False):
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def _add_para(container, text="", align=WD_ALIGN_PARAGRAPH.LEFT,
              bold=False, size=FONT_SIZE, space_before=0, space_after=4):
    p = container.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        _set_font(run, size, bold)
    return p


def _add_run(para, text, bold=False, size=FONT_SIZE):
    run = para.add_run(text)
    _set_font(run, size, bold)
    return run


def _set_cell_bg(cell, hex_color="D9D9D9"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _header_row(table, headers):
    """headers = [(label, width_cm), ...]"""
    row = table.add_row()
    for i, (label, w) in enumerate(headers):
        c = row.cells[i]
        c.width = Cm(w)
        _set_cell_bg(c, "FFFFFF")
        p   = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(label)
        _set_font(run, FONT_SIZE, bold=True)
    return row


def _data_row(table, values):
    """values = [(text, width_cm, align), ...]"""
    row = table.add_row()
    for i, (text, w, align) in enumerate(values):
        c = row.cells[i]
        c.width = Cm(w)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1
        lines = str(text).split("\n")
        for idx, line in enumerate(lines):
            run = p.add_run(line)
            if i == 2:
                run.font.color.rgb = RGBColor(0, 0, 255)
            _set_font(run, Pt(12))
            if idx < len(lines) - 1:
                run.add_break()
    return row


def _signature_block(doc, signer_rank, signer_name, signer_pos):
    _add_para(doc, "", space_before=2, space_after=2)
    _add_para(doc, signer_rank,  WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "",           space_before=2, space_after=2)
    _add_para(doc, signer_name,  WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, signer_pos,   WD_ALIGN_PARAGRAPH.CENTER)


def _set_page_margin(doc, top=2.5, bottom=2.5, left=2.5, right=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


# ── Template 1: ร้อยเวร 60 ─────────────────────────────────────
def build_template1(vals: dict) -> BytesIO:
    doc = Document()
    _set_page_margin(doc)

    _add_para(doc, "คำสั่งงานจราจร สถานีตำรวจภูธรตระการพืชผล",
              WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=FONT_SIZE_TITLE, space_after=2)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _add_run(p, "ที่  ", bold=True)
    _add_run(p, vals.get("ORDER_NO", "{{ORDER_NO}}"))

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    _add_run(p2, "เรื่อง  ", bold=True)
    _add_run(p2, "มอบหมายการปฏิบัติหน้าที่ของร้อยเวรจราจร และสายตรวจจราจร", bold=True)

    _add_para(doc, "………………………….………………..",
              WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=4)

    body1 = doc.add_paragraph()
    body1.paragraph_format.space_after  = Pt(4)
    body1.paragraph_format.first_line_indent = Cm(2.5)
    _add_run(body1,
        f"ตามคำสั่ง สถานีตำรวจภูธรตระการพืชผล ที่ {vals.get('REF_ORDER_NO','{{REF_ORDER_NO}}')} "
        f"ลงวันที่ {vals.get('REF_DATE','{{REF_DATE}}')} "
        "เรื่อง การมอบหมายอำนาจหน้าที่รับผิดชอบและการปฏิบัติราชการในสถานีตำรวจภูธรตระการพืชผล "
        "และอาศัยอำนาจตามคำสั่งสำนักงานตำรวจแห่งชาติ ที่ 537/2555 ลงวันที่ 27 กันยายน 2555 "
        "เรื่อง การกำหนดอำนาจหน้าที่ของตำแหน่งในสถานีตำรวจในด้านงานจราจรและการให้บริการประชาชน"
        "ให้มีประสิทธิภาพ เพื่อให้เกิดประโยชน์สูงสุดในเขตพื้นที่รับผิดชอบความแจ้งอยู่แล้วนั้น")

    body2 = doc.add_paragraph()
    body2.paragraph_format.space_after  = Pt(4)
    body2.paragraph_format.first_line_indent = Cm(2.5)
    _add_run(body2,
        "สถานีตำรวจภูธรตระการพืชผล จึงได้แต่งตั้งเจ้าหน้าที่ตำรวจเพื่อทำหน้าที่ร้อยเวรจราจร "
        "(ร้อยเวร 60) พร้อมด้วยลูกชุดในแต่ละวันให้ออกปฏิบัติ ในเขตพื้นที่รับผิดชอบเพื่ออำนวยความ"
        "สะดวกด้านจราจร และจัดการจราจรเมื่อมีอุบัติเหตุบนท้องถนน และปฏิบัติตามสั่งการของ"
        "ผู้บังคับบัญชาที่มอบหมาย ")
    _add_run(body2, f"ประจำเดือน {vals.get('MONTH_YEAR','{{MONTH_YEAR}}')}", bold=True)
    _add_run(body2, " ดังนี้.-")

    for item in ["1. ชุดปฏิบัติการร้อยเวรจราจร (ร้อยเวร 60) ผนวก ก.",
                 "2. ชุดปฏิบัติการสายตรวจจราจร ผนวก ข."]:
        pi = doc.add_paragraph()
        pi.paragraph_format.left_indent = Cm(8)
        pi.paragraph_format.space_after  = Pt(2)
        _add_run(pi, item)

    sup = doc.add_paragraph()
    sup.paragraph_format.space_before = Pt(4)
    sup.paragraph_format.space_after  = Pt(4)
    sup.paragraph_format.first_line_indent = Cm(2.5)
    _add_run(sup, f"ทั้งนี้ มอบให้ {vals.get('SUPERVISOR_NAME','{{SUPERVISOR_NAME}}')}  "
                  f"{vals.get('SUPERVISOR_POS','{{SUPERVISOR_POS}}')} เป็นผู้ควบคุมการปฏิบัติอย่างใกล้ชิด")

    sp = doc.add_paragraph()
    sp.paragraph_format.left_indent = Cm(8)
    sp.paragraph_format.space_before = Pt(4)
    sp.paragraph_format.space_after  = Pt(2)
    _add_run(sp, f"สั่ง   ณ   วันที่   {vals.get('SIGN_DATE','{{SIGN_DATE}}')}")

    _signature_block(doc,
        vals.get("SIGNER_RANK", "{{SIGNER_RANK}}"),
        vals.get("SIGNER_NAME", "{{SIGNER_NAME}}"),
        vals.get("SIGNER_POS",  "{{SIGNER_POS}}"))

    doc.add_page_break()

    # ── ผนวก ก. ──
    _add_para(doc, "ผนวก ก. : ตารางการปฏิบัติร้อยเวรจราจร (ร้อยเวร 60)",
              WD_ALIGN_PARAGRAPH.LEFT, bold=True, size=FONT_SIZE_TITLE, space_after=2)
    _add_para(doc,
        f"ประกอบ : คำสั่ง งานจราจร สภ.ตระการพืชผล ที่ "
        f"{vals.get('ANNEX_ORDER_NO','{{ANNEX_ORDER_NO}}')} "
        f"ลง {vals.get('ANNEX_SIGN_DATE','{{ANNEX_SIGN_DATE}}')}",
        WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    _add_para(doc, "." * 120, WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

    COL_A = [1.2, 4.5, 8.8, 2.2]
    tblA  = doc.add_table(rows=0, cols=4)
    tblA.style = "Table Grid"
    _header_row(tblA, [
        ("ลำดับ", COL_A[0]),
        ("ยศ  ชื่อ     ชื่อสกุล", COL_A[1]),
        ("ปฏิบัติหน้าที่ร้อยเวรจราจร", COL_A[2]),
        ("หมายเหตุ", COL_A[3]),
    ])
    _data_row(tblA, [
        ("1.", COL_A[0], WD_ALIGN_PARAGRAPH.CENTER),
        (vals.get("ROW_A1_NAME", "{{ROW_A1_NAME}}"), COL_A[1], WD_ALIGN_PARAGRAPH.LEFT),
        (vals.get("ROW_A1_DATES", "{{ROW_A1_DATES}}"), COL_A[2], WD_ALIGN_PARAGRAPH.LEFT),
        ("", COL_A[3], WD_ALIGN_PARAGRAPH.CENTER),
    ])
    _data_row(tblA, [
        ("2.", COL_A[0], WD_ALIGN_PARAGRAPH.CENTER),
        (vals.get("ROW_A2_NAME", "{{ROW_A2_NAME}}"), COL_A[1], WD_ALIGN_PARAGRAPH.LEFT),
        (vals.get("ROW_A2_DATES", "{{ROW_A2_DATES}}"), COL_A[2], WD_ALIGN_PARAGRAPH.LEFT),
        ("", COL_A[3], WD_ALIGN_PARAGRAPH.CENTER),
    ])

    _add_para(doc, "", space_before=4, space_after=2)
    _signature_block(doc,
        vals.get("SIGNER_RANK", "{{SIGNER_RANK}}"),
        vals.get("SIGNER_NAME", "{{SIGNER_NAME}}"),
        vals.get("SIGNER_POS",  "{{SIGNER_POS}}"))

    doc.add_page_break()

    # ── ผนวก ข. ──
    _add_para(doc, "ผนวก ข. : ตารางการปฏิบัติสายตรวจจราจร",
              WD_ALIGN_PARAGRAPH.LEFT, bold=True, size=FONT_SIZE_TITLE, space_after=2)
    _add_para(doc,
        f"ประกอบ : คำสั่ง งานจราจร สภ.ตระการพืชผล ที่ "
        f"{vals.get('ANNEX_ORDER_NO','{{ANNEX_ORDER_NO}}')} "
        f"ลง {vals.get('ANNEX_SIGN_DATE','{{ANNEX_SIGN_DATE}}')}",
        WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    _add_para(doc, "." * 120, WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

    COL_B = [1.2, 4.5, 8.8, 2.2]
    tblB  = doc.add_table(rows=0, cols=4)
    tblB.style = "Table Grid"
    _header_row(tblB, [
        ("ลำดับ", COL_B[0]),
        ("ยศ  ชื่อ     ชื่อสกุล", COL_B[1]),
        ("ปฏิบัติหน้าที่สายตรวจจราจร", COL_B[2]),
        ("หมายเหตุ", COL_B[3]),
    ])

    for row_no, names_key, dates_key in [
        ("1.", "ROW_B1_NAMES", "ROW_B1_DATES"),
        ("2.", "ROW_B2_NAMES", "ROW_B2_DATES"),
    ]:
        row = tblB.add_row()
        aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER]
        for i, (w, align) in enumerate(zip(COL_B, aligns)):
            row.cells[i].width = Cm(w)
            row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = row.cells[i].paragraphs[0]
            p.alignment = align
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
        _set_font(row.cells[0].paragraphs[0].add_run(row_no))
        _set_font(row.cells[1].paragraphs[0].add_run(vals.get(names_key, f"{{{{{names_key}}}}}")))
        _set_font(row.cells[2].paragraphs[0].add_run(vals.get(dates_key, f"{{{{{dates_key}}}}}")))

    _data_row(tblB, [
        ("3.", COL_B[0], WD_ALIGN_PARAGRAPH.CENTER),
        (vals.get("ROW_B3_NAME", "{{ROW_B3_NAME}}"), COL_B[1], WD_ALIGN_PARAGRAPH.LEFT),
        (vals.get("ROW_B3_ROLE", "{{ROW_B3_ROLE}}"), COL_B[2], WD_ALIGN_PARAGRAPH.LEFT),
        ("", COL_B[3], WD_ALIGN_PARAGRAPH.CENTER),
    ])

    _add_para(doc, "", space_before=4, space_after=2)
    _signature_block(doc,
        vals.get("SIGNER_RANK", "{{SIGNER_RANK}}"),
        vals.get("SIGNER_NAME", "{{SIGNER_NAME}}"),
        vals.get("SIGNER_POS",  "{{SIGNER_POS}}"))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# =================================================================
# ── DUTY SPOTS (10 จุดคงที่) ────────────────────────────────────
# =================================================================
DUTY_SPOTS_FIXED = [
    "๑.สี่แยกบ้านดอนใหญ่",
    "๒.ยูเทิร์น ปั้ม ปตท.",
    "๓.สืบสวนหาข่าว(รอตั้งจุดตรวจ)",
    "๔.หน้า รร.มัธยมตระการพืชผล",
    "๕.หน้า รร.อนุบาลน้องหญิง",
    "๖.สืบสวนหาข่าว",
    "๗.สี่แยกธนาคารกรุงเทพ",
    "๘.รร.อนุบาลกุลวิกรานต์,(ตลาดขนส่ง ๑๖.๐๐-๑๗.๓๐ น.)",
    "๙.รร.อนุบาลตระการพืชผล,(ตลาดเทศบาล๑ ๑๖.๐๐-๑๗.๓๐ น.)",
    "๑๐.สืบสวนหาข่าว",
]

TRAFFIC_OFFICERS_FIXED = [
    "ร.ต.ท.พจน์  ปาสาจันทร์ ( ตระการ 614 )",
    "ร.ต.ต.ประจวบ ศรีบุระ ( ตระการ 615 )",
    "ด.ต.ยุทธพงษ์  ชาตแดง ( ตระการ 6101 )",
    "จ.ส.ต.วัชระ จันสุตะ ( ตระการ 6102 )",
    "ส.ต.อ.ธีระวัฒน์  แก่นสาร ( ตระการ 6103 )",
    "ส.ต.อ.ณรงค์ฤทธิ์  เหล่าดี ( ตระการ 6104 )",
    "ส.ต.อ.วิชชากร  วงษ์โท ( ตระการ 6105 )",
    "ส.ต.ท.วิศรุต จันทร์สิงห์ ( ตระการ 6106 )",
    "ส.ต.ท.วชิรวิชญ์ นันทรักษ์ ( ตระการ 6108 )",
    "ส.ต.ต.อดิศร    ศุภนิกร ( ตระการ 6111 )",
]

# ── สีตามจุด (ต้องนิยามก่อนฟังก์ชันที่ใช้) ─────────────────────
SPOT_COLORS = {
    2: RGBColor(0x00, 0x70, 0xC0),   # น้ำเงิน
    5: RGBColor(0xFF, 0x00, 0x00),   # แดง
    9: RGBColor(0xFF, 0x00, 0x00),   # แดง
}


def _get_spot_days(days_list: list, person_idx: int, n_spots: int = 10) -> dict:
    """
    คืน dict {spot_idx(0-based): [วันที่...]}
    - คนที่ person_idx เริ่มจุด person_idx % n_spots
    - วันที่แรกเริ่มที่ days_list[person_idx % n]
    - วนรอบทั้งจุดและวันที่แบบ circular
    """
    spot_days = defaultdict(list)
    n = len(days_list)
    if n == 0:
        return spot_days

    spot_offset = person_idx % n_spots
    day_offset  = person_idx % n

    for i in range(n):
        day_idx  = (day_offset + i) % n
        spot_idx = (spot_offset + i) % n_spots
        spot_days[spot_idx].append(days_list[day_idx])

    return spot_days


def _fill_duty_cell_split(cell_name, cell_dates, spot_days: dict, n_spots: int = 10):
    for p in cell_name.paragraphs:
        p.clear()
    for p in cell_dates.paragraphs:
        p.clear()

    for s in range(n_spots):
        spot_label = DUTY_SPOTS_FIXED[s]
        dates_list = spot_days.get(s, [])
        color      = SPOT_COLORS.get(s, RGBColor(0, 0, 0))

        # ── คอลัมน์ชื่อจุด ──────────────────────────────────
        if s == 0:
            p_name = cell_name.paragraphs[0]
        else:
            p_name = cell_name.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_name.paragraph_format.space_before = Pt(0)
        p_name.paragraph_format.space_after  = Pt(0)
        p_name.paragraph_format.line_spacing = 1
        run_name = p_name.add_run(spot_label)
        run_name.font.color.rgb = color
        _set_font(run_name, Pt(12))

        # ── คอลัมน์วันที่: จับคู่ทีละ 2 ตามลำดับใน dates_list ──
        pairs = []
        if len(dates_list) == 0:
            pairs.append(("", ""))
        else:
            for pi in range(0, len(dates_list), 2):
                l = str(dates_list[pi])
                r = str(dates_list[pi + 1]) if pi + 1 < len(dates_list) else ""
                pairs.append((l, r))

        for pair_i, (left_val, right_val) in enumerate(pairs):
            if s == 0 and pair_i == 0:
                p_date = cell_dates.paragraphs[0]
            else:
                p_date = cell_dates.add_paragraph()
            p_date.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_date.paragraph_format.space_before = Pt(0)
            p_date.paragraph_format.space_after  = Pt(0)
            p_date.paragraph_format.line_spacing = 1

            left_text = f"{left_val}\t" if (left_val and right_val) else left_val
            run_left = p_date.add_run(left_text)
            run_left.font.color.rgb = color
            _set_font(run_left, Pt(12))

            run_right = p_date.add_run(right_val)
            run_right.font.color.rgb = color
            _set_font(run_right, Pt(13))

# ── Template 2: เวรประจำจุดวันธรรมดา / วันหยุด ──────────────────────────
def _build_traffic_doc(vals: dict, holiday: bool = False) -> BytesIO:

    doc = Document()
    _set_page_margin(doc)

    # =========================
    # โลโก้ครุฑ
    # =========================
    try:
        BASE_DIR  = os.path.dirname(os.path.abspath(__file__))
        logo_path = os.path.join(BASE_DIR, "garuda.png")
        if os.path.exists(logo_path):
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(logo_path, width=Inches(0.8))
    except Exception as e:
        print("ERROR LOGO :", e)

    title_sub = (
        "แต่งตั้งเจ้าหน้าที่ตำรวจสายตรวจจราจรทำหน้าที่เวรประจำจุด "
        "ในวันหยุดและวันหยุดนักขัตฤกษ์"
        if holiday else
        "แต่งตั้งเจ้าหน้าที่ตำรวจสายตรวจจราจรทำหน้าที่เวรประจำจุด"
        "ทางแยกทางร่วมและหน้าสถานศึกษาในเขตชุมชน"
    )

    # =========================
    # หัวเอกสาร
    # =========================
    _add_para(doc, "คำสั่งงานจราจร สถานีตำรวจภูธรตระการพืชผล",
              WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=Pt(14), space_after=0)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _add_run(p, "ที่  ", bold=True)
    _add_run(p, vals.get("ORDER_NO", ""))

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    _add_run(p2, "เรื่อง  ", bold=True)
    _add_run(p2, title_sub, bold=True)

    _add_para(doc, "........................................................................",
              WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=4)

    # =========================
    # เนื้อหา
    # =========================
    body = doc.add_paragraph()
    body.paragraph_format.space_after       = Pt(6)
    body.paragraph_format.first_line_indent = Cm(2.5)

    if holiday:
        _add_run(body,
            "เพื่อให้การปฏิบัติหน้าที่ของสายตรวจจราจร สถานีตำรวจภูธรตระการพืชผล "
            "ในการอำนวยความสะดวกการจราจรและแก้ไขปัญหาจราจรในเขตพื้นที่รับผิดชอบ "
            "เป็นไปด้วยความเรียบร้อย จึงมอบให้สายตรวจจราจรอำนวยความสะดวกด้านการจราจร "
            "ในวันหยุดและวันหยุดนักขัตฤกษ์ ")
        _add_run(body, f"ประจำเดือน {vals.get('MONTH_YEAR','')}", bold=True)
        _add_run(body, "  ดังนี้")
    else:
        _add_run(body,
            "เพื่อให้การปฏิบัติหน้าที่ของสายตรวจจราจร สถานีตำรวจภูธรตระการพืชผล "
            "ในการอำนวยความสะดวกการจราจรและแก้ไขปัญหาจราจรในเขตพื้นที่รับผิดชอบ "
            "เป็นไปด้วยความเรียบร้อย จึงมอบให้สายตรวจจราจรอำนวยความสะดวกด้านการจราจร "
            "ในช่วงโมงเร่งด่วนและประจำจุดที่รับผิดชอบ ")
        _add_run(body, f"ประจำเดือน {vals.get('MONTH_YEAR','')}", bold=True)
        _add_run(body,
            f" (ช่วงเช้าเวลา {vals.get('MORNING_TIME','07.00-08.00 น.')} "
            f"และช่วงบ่ายเวลา {vals.get('AFTERNOON_TIME','15.30-16.30 น.')}) ดังนี้")

    # =========================
    # ข้อมูลวันที่และกำลังพล
    # =========================
    base_dates  = vals.get("CUSTOM_DATES", list(range(1, 11)))
    n_spots     = len(DUTY_SPOTS_FIXED)

    # =========================
    # ตาราง 4 คอลัมน์
    # คอลัมน์: ลำดับ | ยศ ชื่อ | เวรประจำจุด+วันที่ | หมายเหตุ
    # =========================
    COL = [1.0, 3.8, 8, 3.0, 1.4]

    tbl = doc.add_table(rows=0, cols=5)
    tbl.style = "Table Grid"
    tbl.autofit = False

    # ── ฟังก์ชันภายใน (indent ถูกต้อง) ──────────────────────
    def add_header():
        return _header_row(tbl, [
            ("ลำดับ",          COL[0]),
            ("ยศ ชื่อ - สกุล", COL[1]),
            ("เวรประจำจุด",    COL[2]),
            ("วันที่ปฏิบัติ",  COL[3]),
            ("หมายเหตุ",       COL[4]),
        ])

    def set_repeat_header(row):
        """กำหนดให้แถวนี้เป็น header ที่ซ้ำทุกหน้า"""
        trPr = row._tr.get_or_add_trPr()
        tblHeader = OxmlElement("w:tblHeader")
        tblHeader.set(qn("w:val"), "1")
        trPr.append(tblHeader)

    # ── เพิ่ม header แถวแรก ───────────────────────────────────
    hdr_row = add_header()
    set_repeat_header(hdr_row)

    repeat_at = {4, 7} if not holiday else {5}

    for i, officer_name in enumerate(TRAFFIC_OFFICERS_FIXED):
        sd  = _get_spot_days(base_dates, i, n_spots)
        row = tbl.add_row()

        # col 0: ลำดับ
        c0 = row.cells[0]; c0.width = Cm(COL[0])
        c0.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after  = Pt(0)
        _set_font(p0.add_run(f"{i+1}."), Pt(12))

        # col 1: ยศ ชื่อ แยก 2 บรรทัด
        c1 = row.cells[1]; c1.width = Cm(COL[1])
        c1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p1_blank = c1.paragraphs[0]
        p1_blank.paragraph_format.space_before = Pt(0)
        p1_blank.paragraph_format.space_after  = Pt(0)
        p1_blank.paragraph_format.line_spacing = 1

        if " ( " in officer_name:
            name_part, code_part = officer_name.split(" ( ", 1)
            code_part = "( " + code_part
        else:
            name_part = officer_name
            code_part = ""

        p1_name = c1.add_paragraph()
        p1_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1_name.paragraph_format.space_before = Pt(0)
        p1_name.paragraph_format.space_after  = Pt(0)
        p1_name.paragraph_format.line_spacing = 1
        _set_font(p1_name.add_run(name_part.strip()), Pt(12))

        if code_part:
            p1_code = c1.add_paragraph()
            p1_code.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p1_code.paragraph_format.space_before = Pt(0)
            p1_code.paragraph_format.space_after  = Pt(0)
            p1_code.paragraph_format.line_spacing = 1
            _set_font(p1_code.add_run(code_part.strip()), Pt(12))

        # col 2 & 3: ชื่อจุด | วันที่
        c2 = row.cells[2]; c2.width = Cm(COL[2])
        c2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        c3 = row.cells[3]; c3.width = Cm(COL[3])
        c3.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _fill_duty_cell_split(c2, c3, sd, n_spots)

        # col 4: หมายเหตุ
        c4 = row.cells[4]; c4.width = Cm(COL[4])
        c4.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p4 = c4.paragraphs[0]
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.paragraph_format.space_before = Pt(0)
        p4.paragraph_format.space_after  = Pt(0)
        _set_font(p4.add_run(""), Pt(12))

        if (i + 1) in repeat_at and (i + 1) < len(TRAFFIC_OFFICERS_FIXED):
            add_header()

    # =========================
    # ผู้ควบคุม + ลายเซ็น
    # =========================
    _add_para(doc, "", space_before=6, space_after=2)

    sup = doc.add_paragraph()
    sup.paragraph_format.space_after       = Pt(4)
    sup.paragraph_format.first_line_indent = Cm(2.5)
    _add_run(sup,
        f"ทั้งนี้ มอบหมายให้ {vals.get('SUPERVISOR_NOTE','')} "
        "เป็นผู้ควบคุมการปฏิบัติอย่างใกล้ชิด")

    sp = doc.add_paragraph()
    sp.paragraph_format.left_indent  = Cm(8)
    sp.paragraph_format.space_before = Pt(4)
    sp.paragraph_format.space_after  = Pt(2)
    _add_run(sp, f"สั่ง   ณ   วันที่   {vals.get('SIGN_DATE','')}")

    _signature_block(doc,
        vals.get("SIGNER_RANK", ""),
        vals.get("SIGNER_NAME", ""),
        vals.get("SIGNER_POS",  ""))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
# =================================================================
# HELPER FUNCTIONS (arrest system)
# =================================================================
def date_th(d):
    if not d:
        return "-"
    months = ["","มกราคม","กุมภาพันธ์","มีนาคม","เมษายน","พฤษภาคม","มิถุนายน",
              "กรกฎาคม","สิงหาคม","กันยายน","ตุลาคม","พฤศจิกายน","ธันวาคม"]
    try:
        return f"{d.day} {months[d.month]} {d.year + 543}"
    except Exception:
        return str(d)


def safe_time(t):
    if isinstance(t, (datetime.time, datetime.datetime)):
        return t.strftime("%H:%M")
    return str(t) if t else "-"


def to_arabic_number(text):
    if text is None:
        return ""
    m = {"๐":"0","๑":"1","๒":"2","๓":"3","๔":"4",
         "๕":"5","๖":"6","๗":"7","๘":"8","๙":"9"}
    return "".join(m.get(c, c) for c in str(text))


def format_thai_id(id_number):
    digits = "".join(filter(str.isdigit, str(id_number or "")))
    if len(digits) != 13:
        return id_number or ""
    return f"{digits[0]}-{digits[1:5]}-{digits[5:10]}-{digits[10:12]}-{digits[12]}"


def clean_geo(text):
    if not text:
        return ""
    text = str(text).strip()
    for w in ["จังหวัด","อำเภอ","ตำบล"]:
        text = text.replace(w, "")
    return text.strip()


def format_amphur(name):
    if not name:
        return ""
    for x in ["อำเภอ","อ.","เขต"]:
        name = str(name).replace(x, "")
    return name.strip()


def format_tambon(name):
    return str(name or "").replace("ตำบล","").strip()


def safe_index(value, options):
    return options.index(value) if value in options else 0


def remove_empty_signature_rows(doc):
    for table in doc.tables:
        rows_to_remove = []
        for row in table.rows:
            if len(row.cells) < 3:
                continue
            center = row.cells[1].text.strip()
            right  = row.cells[2].text.strip()
            if right == "" and "ผู้จับกุม" in center:
                rows_to_remove.append(row)
        for row in rows_to_remove:
            table._tbl.remove(row._tr)


def replace_text(doc, data):
    def smart_replace(paragraph):
        full_text = "".join(run.text for run in paragraph.runs)
        for key, value in data.items():
            if key in full_text:
                full_text = full_text.replace(key, str(value))
        if paragraph.runs:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""
    for p in doc.paragraphs:
        smart_replace(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    smart_replace(p)
    remove_empty_signature_rows(doc)


@st.cache_data
def load_geo():
    try:
        df = pd.read_csv("thai_districts.csv", encoding="utf-8-sig")
    except Exception:
        df = pd.read_csv("thai_districts.csv", encoding="cp874")
    df["ProvinceThai"] = df["ProvinceThai"].apply(clean_geo)
    df["DistrictThai"] = df["DistrictThai"].apply(clean_geo).apply(format_amphur)
    df["TambonThai"]   = df["TambonThai"].apply(clean_geo).apply(format_tambon)
    return df


def get_base64(path):
    if not os.path.exists(path):
        return ""
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode()


VEHICLE_MODELS = {
    "Honda":  ["Wave 110i","Wave 125i","Click 125i","PCX","ADV160"],
    "Yamaha": ["NMAX","Aerox","Grand Filano","Fino"],
    "Toyota": ["Vios","Yaris","Hilux Revo","Fortuner"],
    "Isuzu":  ["D-Max","MU-X"],
    "Nissan": ["Navara","Almera"],
    "Mazda":  ["Mazda2","Mazda3","BT-50"],
}

OFFICERS = [
    "ร.ต.อ.สมคิด เชื้อเวียง","ร.ต.อ.ชัยพร ชอบงาม","ร.ต.ท.พจน์ ปาสาจันทร์",
    "ร.ต.ต.ประจวบ ศรีบุระ","ส.ต.อ.ธีระวัฒน์ แก่นสาร","ส.ต.อ.ณรงค์ฤทธิ์ เหล่าดี",
    "ส.ต.อ.วิชชากร วงษ์โท","ด.ต.ยุทธพงษ์ ชาดแดง","จ.ส.ต.วัชระ จันสุตะ",
    "ส.ต.ท.วชิรวิชญ์ นันทรักษ์","ส.ต.ท.วิศรุต จันทร์สิงห์","ส.ต.ท.เอกพจน์ อินผล",
    "ส.ต.ต.อดิศร ศุภนิกร",
]

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
logo_b64 = st.session_state.settings.get("logo") or get_base64(os.path.join(BASE_DIR, "police_logo.png"))

PLOTLY_LAYOUT = dict(
    paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
    font=dict(family="Sarabun, sans-serif", color="#5A7A99", size=12),
    margin=dict(l=20, r=20, t=44, b=20),
    title_font=dict(size=15, color="#003B6F"),
    colorway=["#00AEEF","#0076B6","#003B6F","#F5A623","#2EC4B6","#E84855"],
    xaxis=dict(gridcolor="rgba(0,174,239,0.08)", linecolor="rgba(0,0,0,0.08)"),
    yaxis=dict(gridcolor="rgba(0,174,239,0.08)", linecolor="rgba(0,0,0,0.08)"),
)

month_names = ["","มกราคม","กุมภาพันธ์","มีนาคม","เมษายน","พฤษภาคม","มิถุนายน",
               "กรกฎาคม","สิงหาคม","กันยายน","ตุลาคม","พฤศจิกายน","ธันวาคม"]

# =================================================================
# LOGIN
# =================================================================
if "password_correct" not in st.session_state:
    if logo_b64:
        st.markdown(
            f'<div style="text-align:center;margin:2rem 0 1rem;">'
            f'<img src="data:image/png;base64,{logo_b64}" width="90" style="border-radius:16px;">'
            f'</div>', unsafe_allow_html=True)
    st.markdown("""
    <div style="text-align:center;margin-bottom:1.5rem;">
        <div style="font-size:28px;font-weight:800;color:var(--ktb-blue);">ยินดีต้อนรับ</div>
        <div style="font-size:15px;color:var(--text-muted);margin-top:4px;">
            ระบบบันทึกการจับกุม งานจราจร<br>
            <strong>สถานีตำรวจภูธรตระการพืชผล</strong>
        </div>
    </div>""", unsafe_allow_html=True)
    _, c, _ = st.columns([1, 2, 1])
    with c:
        u = st.text_input("👤 ชื่อผู้ใช้งาน", key="lu")
        p = st.text_input("🔒 รหัสผ่าน", type="password", key="lp")
        st.write("")
        if st.button("เข้าสู่ระบบ", use_container_width=True):
            ud = USERS.get(u)
            if ud and ud["password"] == p:
                st.session_state["password_correct"] = True
                st.session_state["user_full_name"]   = ud["fullname"]
                st.session_state.page = "dashboard"
                st.rerun()
            else:
                st.error("❌ ชื่อผู้ใช้หรือรหัสผ่านไม่ถูกต้อง")
    st.stop()

if "page" not in st.session_state:
    st.session_state.page = "dashboard"

records     = st.session_state.get("records", [])
doc_records = st.session_state.get("doc_records", [])

# =================================================================
# SIDEBAR
# =================================================================
with st.sidebar:
    if logo_b64:
        st.markdown(
            f'<div style="text-align:center;margin-bottom:8px;">'
            f'<img src="data:image/png;base64,{logo_b64}" width="80" '
            f'style="border-radius:16px;padding:8px;background:rgba(0,174,239,0.15);"></div>',
            unsafe_allow_html=True)

    st.markdown("""
    <div style="text-align:center;font-size:17px;font-weight:800;color:#00AEEF;margin-bottom:14px;line-height:1.35;">
        ระบบบันทึกจับกุม<br>
        <span style="font-size:13px;font-weight:400;color:#7FB3D3;">งานจราจร</span>
    </div>""", unsafe_allow_html=True)

    st.markdown(
        f'<div class="sidebar-profile">'
        f'👮 <strong>{st.session_state.get("user_full_name","เจ้าหน้าที่")}</strong><br>'
        f'🏢 สภ.ตระการพืชผล</div>', unsafe_allow_html=True)

    nav_items = [
        ("🏠", "หน้าแรก",          "dashboard"),
        ("📋", "บันทึกจับกุม",     "form"),
        ("🔎", "ค้นหาเอกสาร",      "search"),
        ("📂", "จัดการเอกสาร",     "documents"),
        ("📅", "ตารางเวร / คำสั่ง", "schedule"),
        ("⚙️", "ตั้งค่า",           "settings"),
    ]
    for icon, label, page_key in nav_items:
        if st.button(f"{icon}  {label}", use_container_width=True, key=f"nav_{page_key}"):
            st.session_state.page = page_key
            st.rerun()

    st.markdown("<hr style='margin:10px 0;border-color:rgba(0,174,239,0.2);'>", unsafe_allow_html=True)
    if st.button("🚪  ออกจากระบบ", use_container_width=True, key="logout_btn"):
        st.session_state.confirm_logout = True

if st.session_state.get("confirm_logout"):
    st.warning("⚠️ ต้องการออกจากระบบ?")
    cy, cn = st.columns(2)
    if cy.button("✅ ยืนยัน", key="confirm_yes"):
        for k in list(st.session_state.keys()):
            del st.session_state[k]
        st.rerun()
    if cn.button("❌ ยกเลิก", key="confirm_no"):
        st.session_state.confirm_logout = False
        st.rerun()


# =================================================================
# DASHBOARD
# =================================================================
if st.session_state.page == "dashboard":
    total  = len(records)
    today  = len([r for r in records if str(r.get("report_date")) == str(datetime.date.today())])
    latest = records[-1].get("name", "-") if records else "-"

    st.markdown('<div class="main-header">🚔&nbsp; POLICE REALTIME DASHBOARD — สภ.ตระการพืชผล</div>', unsafe_allow_html=True)

    kpi_data = [
        ("📁", total,   "คดีทั้งหมด"),
        ("📅", today,   "วันนี้"),
        ("🖨️", total,   "เอกสารทั้งหมด"),
        ("👤", latest,  "ล่าสุด"),
    ]
    cols = st.columns(4)
    for col, (icon, val, label) in zip(cols, kpi_data):
        col.markdown(
            f'<div class="kpi-card"><div class="kpi-icon">{icon}</div>'
            f'<div class="kpi-number">{val}</div>'
            f'<div class="kpi-label">{label}</div></div>', unsafe_allow_html=True)

    st.write("")
    if not records:
        charges = ["เมาแล้วขับ","ไม่มีใบขับขี่","ขับเร็ว","ฝ่าไฟแดง","ไม่สวมหมวก"]
        sample  = [{
            "report_date": datetime.date.today() - datetime.timedelta(days=int(np.random.randint(0,30))),
            "charge": str(np.random.choice(charges)),
            "hour":   int(np.random.randint(0,24)),
        } for _ in range(60)]
        df = pd.DataFrame(sample)
    else:
        rows = []
        for r in records:
            try:
                d = pd.to_datetime(r.get("report_date", datetime.date.today()))
            except Exception:
                d = pd.to_datetime(datetime.date.today())
            rows.append({"report_date": d, "charge": r.get("charge","อื่นๆ"), "hour": int(np.random.randint(0,24))})
        df = pd.DataFrame(rows)

    df["report_date"] = pd.to_datetime(df["report_date"], errors="coerce")
    df["month"]       = df["report_date"].dt.strftime("%m")
    charge_df = df.groupby("charge").size().reset_index(name="count")
    month_df  = df.groupby("month").size().reset_index(name="count")

    g1, g2 = st.columns([1.5, 1])
    with g1:
        fig = px.bar(month_df, x="month", y="count", text_auto=True, title="📈 สถิติคดีรายเดือน",
                     color_discrete_sequence=["#00AEEF"])
        fig.update_layout(**PLOTLY_LAYOUT, height=380)
        fig.update_traces(marker_line_width=0, marker_cornerradius=4)
        st.plotly_chart(fig, use_container_width=True)
    with g2:
        fig2 = px.pie(charge_df, names="charge", values="count", title="🥧 สัดส่วนข้อหา",
                      color_discrete_sequence=["#00AEEF","#0076B6","#003B6F","#F5A623","#2EC4B6"])
        fig2.update_layout(**PLOTLY_LAYOUT, height=380)
        st.plotly_chart(fig2, use_container_width=True)

    DOC_META = {
        "หนังสือราชการ":"📋","บันทึกข้อความ":"📝","คำสั่ง":"📌",
        "ประกาศ":"📢","รายงาน":"📊","หนังสือรับ":"📥","หนังสือส่ง":"📤",
    }
    type_counts = {t: 0 for t in DOC_META}
    for d in doc_records:
        t = d.get("doc_type","")
        if t in type_counts:
            type_counts[t] += 1
    st.markdown("### 📄 ประเภทเอกสาร")
    dc = st.columns(4)
    for idx, (dtype, icon) in enumerate(DOC_META.items()):
        dc[idx % 4].markdown(
            f'<div class="metric-card"><div class="metric-icon">{icon}</div>'
            f'<div class="metric-no">{type_counts[dtype]}</div>'
            f'<div class="metric-text">{dtype}</div></div>', unsafe_allow_html=True)


# =================================================================
# SETTINGS
# =================================================================
if st.session_state.page == "settings":
    st.markdown('<div class="main-header">⚙️&nbsp; ตั้งค่าระบบ</div>', unsafe_allow_html=True)
    s = st.session_state.settings
    tab1, tab2, tab3, tab4 = st.tabs(["🏢 หน่วยงาน","🎨 ธีม","💾 สำรองข้อมูล","⚙️ ขั้นสูง"])
    with tab1:
        c1, c2 = st.columns(2)
        with c1:
            s["station"]   = st.text_input("ชื่อสถานีตำรวจ",  s.get("station","สภ.ตระการพืชผล"))
            s["province"]  = st.text_input("จังหวัด",          s.get("province","อุบลราชธานี"))
            s["commander"] = st.text_input("ชื่อผู้กำกับ",      s.get("commander",""))
        with c2:
            s["deputy"]    = st.text_input("รองผู้กำกับ",       s.get("deputy",""))
            s["inspector"] = st.text_input("สารวัตร",           s.get("inspector",""))
        logo_file = st.file_uploader("อัปโหลดโลโก้", type=["png","jpg","jpeg"])
        if logo_file:
            s["logo"] = base64.b64encode(logo_file.read()).decode()
            st.success("อัปโหลดโลโก้แล้ว")
    with tab2:
        s["theme"] = st.selectbox("โหมดระบบ", ["light","dark"],
                                   index=0 if s.get("theme","light")=="light" else 1)
    with tab3:
        recs = st.session_state.get("records",[])
        st.info(f"ข้อมูลทั้งหมด {len(recs)} รายการ")
        st.download_button("📥 ดาวน์โหลด Backup",
                           data=json.dumps(recs, ensure_ascii=False, indent=2, default=str),
                           file_name="backup_records.json", mime="application/json")
        up = st.file_uploader("📤 นำเข้าข้อมูล", type=["json"], key="restore_file")
        if up:
            try:
                st.session_state.records = json.load(up)
                st.success("นำเข้าข้อมูลสำเร็จ")
            except Exception:
                st.error("ไฟล์ไม่ถูกต้อง")
    with tab4:
        f = st.session_state.get("form",{})
        auto_text = (
            f"ตามวันเวลาที่เกิดเหตุ เจ้าหน้าที่ตำรวจชุดจับกุมขณะปฏิบัติหน้าที่\n"
            f"พบ {f.get('name','')} ขับขี่{f.get('vehicle_type','')} ยี่ห้อ {f.get('vehicle_brand','')}\n"
            f"หมายเลขทะเบียน {f.get('vehicle_plate','')}\n"
            f"จากการตรวจสอบพบ {f.get('seized_item','')}\n"
            f"จึงแจ้งข้อกล่าวหา {f.get('charge','')}"
        )
        behavior = st.text_area("พฤติการณ์จับกุม", value=f.get("behavior_text", auto_text), height=250)
        st.session_state.form["behavior_text"] = behavior
        s["auto_save"]   = st.toggle("💾 บันทึกอัตโนมัติ", value=s.get("auto_save", True))
        s["thai_number"] = st.toggle("🔢 ใช้เลขไทย",       value=s.get("thai_number", False))
        c3, c4 = st.columns(2)
        with c3:
            if st.button("🗑️ ล้างข้อมูลทั้งหมด", use_container_width=True):
                st.session_state.records = []
                st.success("ล้างข้อมูลแล้ว")
        with c4:
            if st.button("🚪 รีเซ็ตระบบ", use_container_width=True):
                for k in list(st.session_state.keys()):
                    del st.session_state[k]
                st.rerun()
    c1, c2 = st.columns(2)
    with c1:
        if st.button("💾 บันทึกการตั้งค่า", use_container_width=True):
            st.session_state.settings = s
            st.success("บันทึกแล้ว")
    with c2:
        if st.button("⬅️ กลับหน้าหลัก", use_container_width=True):
            st.session_state.page = "dashboard"
            st.rerun()


# =================================================================
# SEARCH
# =================================================================
if st.session_state.page == "search":
    st.markdown('<div class="main-header">🔎&nbsp; ค้นหาเอกสาร</div>', unsafe_allow_html=True)
    recs = st.session_state.get("records",[])
    c1, c2 = st.columns([3,1])
    keyword = c1.text_input("ค้นหาชื่อ / เลขบัตร / ทะเบียนรถ")
    mode    = c2.selectbox("ประเภท", ["ทั้งหมด","ชื่อ","เลขบัตร","ทะเบียน"])
    result  = []
    for r in recs:
        name  = str(r.get("name",""))
        cid   = str(r.get("sub_id",""))
        plate = str(r.get("vehicle_plate",""))
        ok    = True
        if keyword:
            kw = keyword.lower()
            if mode == "ชื่อ":        ok = kw in name.lower()
            elif mode == "เลขบัตร":   ok = kw in cid.lower()
            elif mode == "ทะเบียน":   ok = kw in plate.lower()
            else:                     ok = kw in name.lower() or kw in cid.lower() or kw in plate.lower()
        if ok:
            result.append(r)
    st.info(f"พบ {len(result)} รายการ")
    for i, r in enumerate(result):
        with st.expander(f"📄 {r.get('name','-')} | {r.get('vehicle_plate','-')}"):
            st.write(f"เลขบัตร: {r.get('sub_id','-')}")
            st.write(f"วันที่: {r.get('report_date','-')}")
            st.write(f"ข้อหา: {r.get('charge','-')}")
            c1, c2 = st.columns(2)
            if c1.button("🗑️ ลบ", key=f"del_{i}"):
                st.session_state.records.remove(r)
                st.rerun()
            if c2.button("📋 โหลด", key=f"load_{i}"):
                st.session_state.form = r
                st.session_state.page = "form"
                st.session_state.tab  = 3
                st.rerun()
    if st.button("⬅️ กลับหน้าหลัก", use_container_width=True):
        st.session_state.page = "dashboard"
        st.rerun()


# =================================================================
# DOCUMENTS
# =================================================================
if st.session_state.page == "documents":
    if "doc_records" not in st.session_state:
        st.session_state.doc_records = []
    if "doc_running_numbers" not in st.session_state:
        st.session_state.doc_running_numbers = {}

    DOC_TYPES = {
        "หนังสือราชการ": {"prefix":"ที่",    "icon":"📋"},
        "บันทึกข้อความ": {"prefix":"บันทึก", "icon":"📝"},
        "คำสั่ง":         {"prefix":"คำสั่ง", "icon":"📌"},
        "ประกาศ":         {"prefix":"ประกาศ", "icon":"📢"},
        "รายงาน":         {"prefix":"รายงาน", "icon":"📊"},
        "หนังสือรับ":     {"prefix":"รับที่",  "icon":"📥"},
        "หนังสือส่ง":     {"prefix":"ส่งที่",  "icon":"📤"},
    }

    def get_next_doc_number(doc_type, year_th):
        key = f"{doc_type}_{year_th}"
        n   = st.session_state.doc_running_numbers.get(key, 0) + 1
        st.session_state.doc_running_numbers[key] = n
        return f"0018.13/{n:04d}/{year_th}"

    def get_preview_number(doc_type, year_th):
        key = f"{doc_type}_{year_th}"
        n   = st.session_state.doc_running_numbers.get(key, 0) + 1
        return f"0018.13/{n:04d}/{year_th}"

    st.markdown('<div class="main-header">📂&nbsp; ระบบจัดการเอกสารราชการ</div>', unsafe_allow_html=True)
    sub_page = st.session_state.get("doc_sub_page", "create")

    if sub_page == "create":
        if st.session_state.get("doc_type_show_picker", True):
            st.markdown("## 📂 เลือกประเภทเอกสาร")
            cols = st.columns(4)
            for idx, (dtype, meta) in enumerate(DOC_TYPES.items()):
                col = cols[idx % 4]
                if col.button(f"{meta['icon']}  {dtype}", use_container_width=True, key=f"pick_doc_{dtype}"):
                    st.session_state["doc_type_preselect"]   = dtype
                    st.session_state["doc_type_show_picker"] = False
                    st.rerun()
        else:
            preselect = st.session_state.get("doc_type_preselect", "หนังสือราชการ")
            meta      = DOC_TYPES.get(preselect, {"icon":"📋"})
            year_th   = datetime.date.today().year + 543

            if st.button("⬅️ เปลี่ยนประเภทเอกสาร", key="btn_change_type"):
                st.session_state["doc_type_show_picker"] = True
                st.rerun()

            st.markdown(f"## {meta['icon']} {preselect}")
            ca, cb = st.columns([2,1])
            with ca:
                auto_no = st.toggle("ออกเลขอัตโนมัติ", value=True, key="tog_auto_no")
                if auto_no:
                    doc_no = get_preview_number(preselect, year_th)
                    st.info(f"เลขที่จะออก: **{doc_no}**")
                else:
                    doc_no = st.text_input("เลขที่เอกสาร", key="doc_manual_no")
            with cb:
                doc_date    = st.date_input("วันที่เอกสาร", datetime.date.today(), key="doc_date")
                doc_urgency = st.selectbox("ชั้นความเร็ว", ["ปกติ","ด่วน","ด่วนมาก","ด่วนที่สุด"])

            st.write("---")
            cc, cd = st.columns(2)
            with cc:
                doc_to   = st.text_input("เรียน / ถึง", key="doc_to")
                doc_from = st.text_input("จาก", value=st.session_state.get("user_full_name",""), key="doc_from")
                doc_ref  = st.text_input("อ้างถึง (ถ้ามี)", key="doc_ref")
            with cd:
                doc_subject = st.text_input("เรื่อง", key="doc_subject")
                doc_attach  = st.text_input("สิ่งที่ส่งมาด้วย", key="doc_attach")

            doc_body = st.text_area("เนื้อหา", height=200, key="doc_body")
            st.write("---")
            ce, cf, cg = st.columns(3)
            signer     = ce.selectbox("ผู้ลงนาม", OFFICERS, key="doc_signer")
            signer_pos = cf.text_input("ตำแหน่ง", key="doc_signer_pos")
            doc_dept   = cg.text_input("หน่วยงาน", value="สภ.ตระการพืชผล", key="doc_dept")
            st.write("---")
            b1, b2 = st.columns(2)
            with b1:
                if st.button("💾 บันทึกเอกสาร", use_container_width=True, key="btn_save_doc"):
                    if not doc_subject:
                        st.warning("⚠️ กรุณากรอกเรื่อง")
                    else:
                        final_no = get_next_doc_number(preselect, year_th) if auto_no else doc_no
                        st.session_state.doc_records.append({
                            "doc_number": final_no, "doc_type": preselect,
                            "doc_date": str(doc_date), "doc_subject": doc_subject,
                            "doc_to": doc_to, "doc_from": doc_from, "doc_ref": doc_ref,
                            "doc_attach": doc_attach, "doc_body": doc_body,
                            "doc_urgency": doc_urgency, "signer": signer,
                            "signer_pos": signer_pos, "doc_dept": doc_dept,
                            "created_by": st.session_state.get("user_full_name","-"),
                            "created_at": str(datetime.datetime.now()),
                            "status": "หนังสือรับ" if preselect == "หนังสือรับ" else "หนังสือส่ง",
                        })
                        st.success(f"✅ บันทึกแล้ว เลขที่ {final_no}")

    if st.button("⬅️ กลับหน้าหลัก", use_container_width=True, key="btn_back_docs"):
        st.session_state.page = "dashboard"
        st.rerun()


# =================================================================
# SCHEDULE
# =================================================================
if st.session_state.page == "schedule":
    st.markdown('<div class="main-header">📅&nbsp; ตารางเวร &amp; คำสั่งจราจร</div>', unsafe_allow_html=True)

    sch_tab1, sch_tab2 = st.tabs(["📋 ตารางเวรรายเดือน", "📄 สร้างคำสั่ง Word"])

    # ── แท็บ 1: ตารางเวรรายเดือน ──────────────────────────────────
    with sch_tab1:
        SHIFT_TYPES = {
            "เวรกลางวัน (06:00-18:00)": "🌤️",
            "เวรกลางคืน (18:00-06:00)": "🌙",
            "วันหยุด": "🏖️",
            "ลา": "📝",
        }

        col_m, col_y = st.columns([2, 2])
        now       = datetime.date.today()
        sel_month = col_m.selectbox("เดือน", list(range(1,13)),
                                    format_func=lambda x: month_names[x],
                                    index=now.month - 1, key="sch_month")
        sel_year  = col_y.number_input("ปี พ.ศ.", value=now.year + 543,
                                       min_value=2560, max_value=2580, key="sch_year")
        sel_year_ad = sel_year - 543

        today_key    = str(now)
        today_shifts = {
            o: v for o, v in st.session_state.schedule_data.items()
            if isinstance(v, dict) and v.get(today_key)
        }
        if today_shifts:
            st.markdown("### 🔔 เวรวันนี้")
            for officer, days in today_shifts.items():
                shift = days.get(today_key, "")
                icon  = SHIFT_TYPES.get(shift, "")
                st.info(f"{icon} **{officer}** — {shift}")
        else:
            st.info(f"🔔 วันนี้ ({date_th(now)}) ยังไม่มีข้อมูลเวร")

        st.write("---")
        st.markdown("### ✏️ กำหนดเวร")
        col_a, col_b, col_c, col_d = st.columns([2, 2, 2, 1])
        sch_officer = col_a.selectbox("เลือกเจ้าหน้าที่", OFFICERS, key="sch_officer")
        sch_date    = col_b.date_input("วันที่", value=now, key="sch_date")
        sch_shift   = col_c.selectbox("ประเภทเวร", list(SHIFT_TYPES.keys()), key="sch_shift")

        if col_d.button("✅ บันทึก", use_container_width=True, key="sch_save"):
            if sch_officer not in st.session_state.schedule_data:
                st.session_state.schedule_data[sch_officer] = {}
            st.session_state.schedule_data[sch_officer][str(sch_date)] = sch_shift
            st.success(f"บันทึกเวร {sch_officer} วันที่ {date_th(sch_date)} แล้ว")

        st.write("---")
        st.markdown(f"### 📋 ตารางเวรเดือน{month_names[sel_month]} {sel_year}")

        _, num_days    = calendar.monthrange(sel_year_ad, sel_month)
        days_in_month  = [datetime.date(sel_year_ad, sel_month, d) for d in range(1, num_days + 1)]

        rows = []
        for officer in OFFICERS:
            row          = {"เจ้าหน้าที่": officer}
            officer_data = st.session_state.schedule_data.get(officer, {})
            for d in days_in_month:
                shift     = officer_data.get(str(d), "")
                row[str(d.day)] = SHIFT_TYPES.get(shift, "") if shift else "-"
            rows.append(row)

        df_sch = pd.DataFrame(rows)
        st.dataframe(df_sch, use_container_width=True, hide_index=True)

        st.write("---")
        st.markdown("### 📤 Export ตารางเวร")
        ex1, ex2 = st.columns(2)
        with ex1:
            csv_rows = []
            for officer in OFFICERS:
                officer_data = st.session_state.schedule_data.get(officer, {})
                for d in days_in_month:
                    shift = officer_data.get(str(d), "")
                    if shift:
                        csv_rows.append({"เจ้าหน้าที่": officer, "วันที่": date_th(d), "ประเภทเวร": shift})
            if csv_rows:
                df_export = pd.DataFrame(csv_rows)
                csv_bytes = df_export.to_csv(index=False).encode("utf-8-sig")
                st.download_button("📥 Export CSV", data=csv_bytes,
                                   file_name=f"ตารางเวร_{month_names[sel_month]}_{sel_year}.csv",
                                   mime="text/csv", key="sch_export_csv")
            else:
                st.warning("ยังไม่มีข้อมูลในเดือนนี้")
        with ex2:
            if st.button("📄 Export Word ตารางเวร", use_container_width=True, key="sch_export_word"):
                try:
                    wdoc = Document()
                    wdoc.add_heading(f"ตารางเวรเดือน{month_names[sel_month]} {sel_year}", level=1)
                    table = wdoc.add_table(rows=1, cols=num_days + 1)
                    table.style = "Table Grid"
                    hdr = table.rows[0].cells
                    hdr[0].text = "เจ้าหน้าที่"
                    for i, d in enumerate(days_in_month):
                        hdr[i+1].text = str(d.day)
                    for officer in OFFICERS:
                        officer_data = st.session_state.schedule_data.get(officer, {})
                        row_cells = table.add_row().cells
                        row_cells[0].text = officer
                        for i, d in enumerate(days_in_month):
                            shift = officer_data.get(str(d), "")
                            row_cells[i+1].text = SHIFT_TYPES.get(shift, "") if shift else "-"
                    buf = BytesIO()
                    wdoc.save(buf)
                    buf.seek(0)
                    st.download_button("📥 ดาวน์โหลด Word",
                                       data=buf,
                                       file_name=f"ตารางเวร_{month_names[sel_month]}_{sel_year}.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                       key="sch_dl_word")
                except Exception as e:
                    st.error(f"❌ {e}")

    # ── แท็บ 2: สร้างคำสั่ง Word ──────────────────────────────────
    with sch_tab2:
        st.markdown("## 📄 สร้างคำสั่งงานจราจร สภ.ตระการพืชผล")
        st.info("เลือกประเภทคำสั่ง กรอกข้อมูล แล้วดาวน์โหลดเป็น Word (.docx)")

        order_type = st.radio(
            "ประเภทคำสั่ง",
            ["คำสั่งร้อยเวร 60 (พร้อมผนวก ก, ข)", "คำสั่งเวรประจำจุดวันธรรมดา", "คำสั่งเวรประจำจุดวันหยุด"],
            horizontal=True,
            key="order_type_radio"
        )

        st.markdown("### 🖊️ ข้อมูลทั่วไป")
        c1, c2, c3 = st.columns(3)
        order_no   = c1.text_input("เลขที่คำสั่ง เช่น 18/2569", key="ord_no")
        month_year = c2.text_input(
            "เดือน ปี เช่น พฤษภาคม 2569",
            value=f"{month_names[datetime.date.today().month]} {datetime.date.today().year + 543}",
            key="ord_month"
        )
        sign_date  = c3.text_input(
            "วันที่ออกคำสั่ง เช่น 29 เมษายน 2569",
            value=date_th(datetime.date.today()),
            key="ord_sign_date"
        )

        c4, c5, c6 = st.columns(3)
        signer_rank = c4.text_input("ยศผู้สั่ง", value="พันตำรวจตรี", key="ord_signer_rank")
        signer_name = c5.text_input("ชื่อผู้สั่ง (ใส่วงเล็บ)", value="( รัฐ   โสมสุพรรณ )", key="ord_signer_name")
        signer_pos  = c6.text_input(
            "ตำแหน่งผู้สั่ง",
            value="สารวัตรป้องกันปราบปราม สถานีตำรวจภูธรตระการพืชผล",
            key="ord_signer_pos"
        )

        vals = {
            "ORDER_NO":    order_no,
            "MONTH_YEAR":  month_year,
            "SIGN_DATE":   sign_date,
            "SIGNER_RANK": signer_rank,
            "SIGNER_NAME": signer_name,
            "SIGNER_POS":  signer_pos,
        }

        # ── ร้อยเวร 60 ────────────────────────────────────────────
        if order_type == "คำสั่งร้อยเวร 60 (พร้อมผนวก ก, ข)":
            st.markdown("### 📎 ผนวก ก: ร้อยเวร 60")
            ca, cb = st.columns(2)
            ref_order = ca.text_input("เลขคำสั่งอ้างอิง", key="ref_ord")
            ref_date  = cb.text_input("วันที่คำสั่งอ้างอิง", key="ref_date")
            annex_no  = ca.text_input("เลขที่ผนวก", key="annex_no")
            annex_dt  = cb.text_input("วันที่ผนวก", value=sign_date, key="annex_dt")

            st.markdown("**แถวที่ 1 (ผนวก ก)**")
            row_a1_name  = st.text_input("ยศ ชื่อ สกุล", key="a1_name")
            row_a1_dates = st.text_area("วันที่ปฏิบัติ (บรรยาย)", key="a1_dates", height=80)

            st.markdown("**แถวที่ 2 (ผนวก ก)**")
            row_a2_name  = st.text_input("ยศ ชื่อ สกุล ", key="a2_name")
            row_a2_dates = st.text_area("วันที่ปฏิบัติ (บรรยาย) ", key="a2_dates", height=80)

            st.markdown("### 📎 ผนวก ข: สายตรวจ")
            st.markdown("**ชุดที่ 1**")
            row_b1_names = st.text_area("ชื่อ (หลายบรรทัด คั่นด้วย Enter)", key="b1_names", height=80)
            row_b1_dates = st.text_area("วันที่ปฏิบัติ", key="b1_dates", height=60)

            st.markdown("**ชุดที่ 2**")
            row_b2_names = st.text_area("ชื่อ (หลายบรรทัด คั่นด้วย Enter) ", key="b2_names", height=80)
            row_b2_dates = st.text_area("วันที่ปฏิบัติ ", key="b2_dates", height=60)

            st.markdown("**แถวธุรการ (แถวที่ 3)**")
            row_b3_name = st.text_input("ชื่อ", key="b3_name")
            row_b3_role = st.text_input("บทบาท / หน้าที่", key="b3_role")

            vals.update({
                "REF_ORDER_NO":    ref_order,
                "REF_DATE":        ref_date,
                "ANNEX_ORDER_NO":  annex_no,
                "ANNEX_SIGN_DATE": annex_dt,
                "ROW_A1_NAME":     row_a1_name,
                "ROW_A1_DATES":    row_a1_dates,
                "ROW_A2_NAME":     row_a2_name,
                "ROW_A2_DATES":    row_a2_dates,
                "ROW_B1_NAMES":    row_b1_names,
                "ROW_B1_DATES":    row_b1_dates,
                "ROW_B2_NAMES":    row_b2_names,
                "ROW_B2_DATES":    row_b2_dates,
                "ROW_B3_NAME":     row_b3_name,
                "ROW_B3_ROLE":     row_b3_role,
            })

            if st.button("💾 สร้างคำสั่งร้อยเวร 60", use_container_width=True, key="gen_t1"):
                with st.spinner("กำลังสร้างเอกสาร..."):
                    buf = build_template1(vals)
                    st.download_button(
                        "📥 ดาวน์โหลดคำสั่งร้อยเวร 60",
                        data=buf,
                        file_name=f"คำสั่งร้อยเวร60_{order_no or 'draft'}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="dl_t1"
                    )

        # ── เวรประจำจุดวันธรรมดา / วันหยุด ───────────────────────
        elif order_type in ("คำสั่งเวรประจำจุดวันธรรมดา", "คำสั่งเวรประจำจุดวันหยุด"):
            is_holiday = (order_type == "คำสั่งเวรประจำจุดวันหยุด")

            st.markdown("### ⚙️ ตั้งค่า")

            if not is_holiday:
                cm1, cm2 = st.columns(2)
                morning_time   = cm1.text_input("เวลาเช้า เช่น 07.00-08.00 น.",   value="07.00-08.00 น.",  key="morning_t")
                afternoon_time = cm2.text_input("เวลาบ่าย เช่น 15.30-16.30 น.", value="15.30-16.30 น.", key="afternoon_t")
                vals["MORNING_TIME"]   = morning_time
                vals["AFTERNOON_TIME"] = afternoon_time

            vals["SUPERVISOR_NOTE"] = st.text_input(
                "ข้อความผู้ควบคุม",
                value="ร้อยเวรจราจร ( 60 ) ประจำผลัด สถานีตำรวจภูธรตระการพืชผล",
                key="sup_note"
            )

            st.markdown("### 📅 กำหนดชุดวันที่ปฏิบัติ")
            st.caption("กรอกวันที่คั่นด้วยช่องว่างหรือลูกน้ำ — คนที่ 1 เริ่มจุดที่ 1 วันแรก, คนถัดไปหมุนเวียนอัตโนมัติ")

            custom_dates_input = st.text_input(
                "ลำดับวันที่",
                value="18 1 19 5 20 6 21 7 22 8 25 12 26 13 27 14 28 15 29",
                key="custom_dates_input"
            )

            # แปลงเป็น list[int]
            import re as _re
            raw_tokens = _re.split(r"[\s,]+", custom_dates_input.strip())
            custom_dates_list = []
            for tok in raw_tokens:
                try:
                    custom_dates_list.append(int(tok))
                except ValueError:
                    pass

            if custom_dates_list:
                st.info(f"✅ พบ {len(custom_dates_list)} วัน: {' → '.join(str(d) for d in custom_dates_list)}")
            else:
                st.warning("⚠️ กรุณากรอกวันที่ให้ถูกต้อง")

            vals["CUSTOM_DATES"] = custom_dates_list

            # ── Preview ตาราง ─────────────────────────────────────
            st.markdown("### 👮 ตาราง Preview (จุด + วันที่ต่อนาย)")

            if custom_dates_list:
                preview_rows = []
                for idx, officer_name in enumerate(TRAFFIC_OFFICERS_FIXED):
                    sd = _get_spot_days(custom_dates_list, idx)
                    for s in range(len(DUTY_SPOTS_FIXED)):
                        dates_str = "\n".join(str(d) for d in sd.get(s, []))
                        preview_rows.append({
                            "นาย": f"{idx+1}. {officer_name[:20]}",
                            "จุด": DUTY_SPOTS_FIXED[s],
                            "วันที่": dates_str,
                        })
                st.dataframe(preview_rows, use_container_width=True, hide_index=True)

            file_label = "วันหยุด" if is_holiday else "วันธรรมดา"

            if st.button(f"💾 สร้างคำสั่งเวร{file_label}", use_container_width=True, key="gen_t23"):
                if not custom_dates_list:
                    st.warning("⚠️ กรุณากรอกวันที่ก่อน")
                else:
                    with st.spinner("กำลังสร้างเอกสาร..."):
                        try:
                            buf = _build_traffic_doc(vals, holiday=is_holiday)
                            st.download_button(
                                f"📥 ดาวน์โหลดคำสั่งเวร{file_label}",
                                data=buf,
                                file_name=f"คำสั่งเวร{file_label}_{order_no or 'draft'}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="dl_t23"
                            )
                        except Exception as e:
                            st.error(f"❌ {e}")
                            st.code(_tb.format_exc())
    st.write("---")
    if st.button("⬅️ กลับหน้าหลัก", use_container_width=True, key="sch_back"):
        st.session_state.page = "dashboard"
        st.rerun()


# =================================================================
# FORM PAGE
# =================================================================
if st.session_state.page != "form":
    st.stop()

st.markdown('<div class="main-header">📋&nbsp; บันทึกการจับกุม</div>', unsafe_allow_html=True)

tab_labels = ["📍 ข้อมูลบันทึก","👮 เจ้าหน้าที่","👤 ผู้ต้องหา","📝 รายละเอียด"]
n1, n2, n3, n4 = st.columns(4)
for i, (col, lbl) in enumerate(zip([n1,n2,n3,n4], tab_labels)):
    with col:
        if st.button(lbl, key=f"nav_tab_{i}", use_container_width=True):
            st.session_state.tab = i
            st.rerun()
st.write("---")


# ======================TAB 1: ข้อมูลบันทึก ====================
# ===========================================================
if st.session_state.tab == 0:
    st.write("")

    col1, col2 = st.columns(2)

    with col1:
        # ================= สถานที่ทำบันทึก (NEW) =================
        record_loc_options = [
            "สภ.ตระการพืชผล ภ.จว.อุบลราชธานี",
            "สภ.เมืองอุบลราชธานี ภ.จว.อุบลราชธานี",
            "สภ.วารินชำราบ ภ.จว.อุบลราชธานี",
            "อื่นๆ"
        ]

        default_loc = st.session_state.form.get(
            "record_loc",
            "สภ.ตระการพืชผล ภ.จว.อุบลราชธานี"
        )

        record_loc_sel = st.selectbox(
            "สถานที่ทำบันทึก",
            record_loc_options,
            index=record_loc_options.index(default_loc) if default_loc in record_loc_options else 0
        )

        # 👉 ถ้าเลือกอื่นๆ ให้เด้ง input
        if record_loc_sel == "อื่นๆ":
            record_loc = st.text_input(
                "ระบุสถานที่ทำบันทึก",
                st.session_state.form.get("record_loc_custom", "")
            )
            st.session_state.form["record_loc_custom"] = record_loc
        else:
            record_loc = record_loc_sel

        # ================= วันที่จับกุม =================
        incident_date = st.date_input(
            "วันที่จับกุม",
            value=st.session_state.form.get('incident_date', datetime.date.today())
        )

        # ================= เวลาจับกุม =================
        arrest_time = st.text_input(
            "เวลาจับกุม (เช่น 10:30)",
            value=str(st.session_state.form.get('arrest_time', "10:00"))
        )

    with col2:
        # ================= วันที่ลงบันทึก =================
        report_date = st.date_input(
            "วันที่ลงบันทึก",
            value=st.session_state.form.get('report_date', datetime.date.today())
        )

        # ================= เวลาลงบันทึก =================
        tz_thai = pytz.timezone("Asia/Bangkok")
        current_time = datetime.datetime.now(tz_thai).replace(second=0, microsecond=0)

        # ล็อคเวลาครั้งแรก
        if "report_datetime_locked" not in st.session_state:
            st.session_state["report_datetime_locked"] = current_time

        locked_datetime = st.session_state["report_datetime_locked"]

        # แยกเวลาไว้ใช้งาน
        report_time = locked_datetime.time()

        st.text_input(
            "เวลาลงบันทึก",
            value=report_time.strftime("%H:%M"),
            disabled=True,
            key="report_time_display"
        )

    # ================= SAVE SESSION =================
    st.session_state.form.update({
        "record_loc": record_loc,
        "incident_date": incident_date,
        "arrest_time": arrest_time,
        "report_date": report_date,

        # เพิ่มบรรทัดนี้
        "report_time": report_time.strftime("%H:%M"),
    })
    # ================= NEXT BUTTON =================
    if st.button("ถัดไป ➔", use_container_width=True, type="primary"):
        st.session_state.tab = 1
        st.rerun()

# --- TAB 2: ผู้บังคับบัญชาและชุดจับกุม ---
if st.session_state.tab == 1:
    st.write("")

    # ================= ผู้บังคับบัญชา =================
    commander_1 = st.text_input(
        "ผกก.",
        st.session_state.form.get(
            'commander_1',
            "พ.ต.อ.เกรียงศักดิ์ ปัชโชโต ผกก.สภ.ตระการพืชผล"
        )
    )

    commander_2 = st.text_input(
        "รอง ผกก.",
        st.session_state.form.get(
            'commander_2',
            "พ.ต.ท.ยุทธนา กิติชัยชนานนท์ รอง ผกก.ป.สภ.ตระการพืชผล"
        )
    )

    commander_3 = st.text_input(
        "สวป.",
        st.session_state.form.get(
            'commander_3',
            "พ.ต.ต.รัฐ โสมสุพรรณ สวป.สภ.ตระการพืชผล"
        )
    )

    st.write("---")

    # =====================================================
    # 🔥 เลือกชุดจับกุมแบบใหม่
    # =====================================================
    st.markdown("### 👮‍♂️ เลือกชุดจับกุม")

    mode = st.radio(
        "เลือกรูปแบบชุด",
        ["ชุดที่ 1", "ชุดที่ 2", "เลือกชุดเอง"],
        horizontal=True
    )

    sub_team = [
        "ร.ต.ท.พจน์ ปาสาจันทร์",
        "ร.ต.ต.ประจวบ ศรีบุระ",
        "ส.ต.อ.ธีระวัฒน์ แก่นสาร",
        "ส.ต.อ.ณรงค์ฤทธิ์ เหล่าดี",
        "ส.ต.อ.วิชชากร วงษ์โท",
        "ด.ต.ยุทธพงษ์ ชาดแดง",
        "จ.ส.ต.วัชระ จันสุตะ",
        "ส.ต.ท.วชิรวิชญ์ นันทรักษ์",
        "ส.ต.ท.วิศรุต จันทร์สิงห์",
        "ส.ต.ท.เอกพจน์ อินผล",
        "ส.ต.ต.อดิศร ศุภนิกร"
    ]

    clean_selected = []

    # ================= ชุดที่ 1 =================
    if mode == "ชุดที่ 1":

        clean_selected = [
            "ร.ต.อ.สมคิด เชื้อเวียง",
            "ร.ต.อ.ชัยพร ชอบงาม",
            *sub_team
        ]

    # ================= ชุดที่ 2 =================
    elif mode == "ชุดที่ 2":

        clean_selected = [
            "ร.ต.อ.ชัยพร ชอบงาม",
            "ร.ต.อ.สมคิด เชื้อเวียง",
            *sub_team
        ]

    # ================= เลือกเอง =================
    else:

        leader = st.selectbox(
            "เลือกหัวหน้าชุด",
            [
                "ร.ต.อ.สมคิด เชื้อเวียง",
                "ร.ต.อ.ชัยพร ชอบงาม"
            ]
        )

        st.markdown("#### 👥 เลือกลูกชุด")

        cols = st.columns(2)
        selected_sub = []

        # 🔥 ใช้ OFFICERS ทั้งหมด (ยกเว้นคนที่เป็นหัวหน้า)
        custom_team = [x for x in OFFICERS if x != leader]

        for i, officer in enumerate(custom_team):
            col = cols[i % 2]

            if col.checkbox(officer, key=f"chk_custom_{i}"):
                selected_sub.append(officer)

        clean_selected = [leader] + selected_sub
        
    # ================= บันทึกลง session =================
    st.session_state.form.update({
        "commander_1": commander_1,
        "commander_2": commander_2,
        "commander_3": commander_3,
        "selected": clean_selected
    })

    # ================= แสดงผล =================
    st.caption(f"เลือกแล้ว {len(clean_selected)} นาย")

    st.text_area(
        "รายชื่อชุดจับกุม",
        value="\n".join(clean_selected),
        height=260
    )

    # ================= ปุ่ม =================
    c1, c2 = st.columns(2)

    with c1:
        if st.button("⬅️ ย้อนกลับ", use_container_width=True):
            st.session_state.tab = 0
            st.rerun()

    with c2:
        if st.button("ถัดไป ➔", use_container_width=True):
            st.session_state.tab = 2
            st.rerun()

# ================= TAB 3 (Index 2): ข้อมูลผู้ต้องหา =================
if st.session_state.tab == 2:
    st.write("")
    # ================= BASIC INFO =================
    col_fname, col_lname, col_id, col_age = st.columns([2,2,2,1])

    first_name = col_fname.text_input(
        "ชื่อ", 
        value=st.session_state.form.get("first_name", ""),  # ← เพิ่ม
        key="input_fname_t3"
    )

    last_name = col_lname.text_input(
        "นามสกุล",
        value=st.session_state.form.get("last_name", ""),   # ← เพิ่ม
        key="input_lname_t3"
    )

    name = f"{first_name} {last_name}"

    sub_id = col_id.text_input(
        "เลขบัตรประชาชน",
        st.session_state.form.get('sub_id', ''),
        key="input_subid_t3"
    )

    age = col_age.number_input(
        "อายุ",
        0, 120,
        st.session_state.form.get('age', 0),
        key="input_age_t3"
    )

    st.write("---")
    st.caption("🏠 ข้อมูลที่อยู่")

    # ================= GEO DATA =================
    df_geo = load_geo()
    # 🔥 safe index กันพัง selectbox
    def safe_index(value, options):
        return options.index(value) if value in options else 0


    # ================= ADDRESS =================
    col_h, col_p = st.columns([1,2])

    house = col_h.text_input(
        "บ้านเลขที่/หมู่",
        st.session_state.form.get('house',''),
        key="input_house_t3"
    )

    provinces = sorted(df_geo['ProvinceThai'].dropna().unique())
    def_prov = st.session_state.form.get('province', provinces[0])

    province = col_p.selectbox(
        "จังหวัด",
        provinces,
        index=safe_index(def_prov, provinces),
        key="sel_prov_t3"
    )

    # ================= DISTRICT =================
    col_a, col_t = st.columns(2)

    districts = sorted(
        df_geo[df_geo['ProvinceThai'] == province]['DistrictThai'].dropna().unique()
    )

    def_amphur = st.session_state.form.get('amphur', '')

    amphur = col_a.selectbox(
        "อำเภอ/เขต",
        districts,
        index=safe_index(def_amphur, districts),
        key="sel_amphur_t3"
    )

    # ================= TAMBON =================
    sub_districts = sorted(
        df_geo[
            (df_geo['ProvinceThai'] == province) &
            (df_geo['DistrictThai'] == amphur)
        ]['TambonThai'].dropna().unique()
    )

    def_tumbon = st.session_state.form.get('tumbon', '')

    tumbon = col_t.selectbox(
        "ตำบล/แขวง",
        sub_districts,
        index=safe_index(def_tumbon, sub_districts),
        key="sel_tumbon_t3"
    )

    # ================= VEHICLE =================
    st.write("---")
    st.caption("🚗 ข้อมูลยานพาหนะ")

    # 🔥 โหลดค่าล่าสุด (ต้องอยู่ก่อน selectbox)
    if "last_vehicle_brand" in st.session_state:
        if not st.session_state.form.get("vehicle_brand"):
            st.session_state.form["vehicle_brand"] = st.session_state["last_vehicle_brand"]

    if "last_vehicle_model" in st.session_state:
        if not st.session_state.form.get("vehicle_model"):
            st.session_state.form["vehicle_model"] = st.session_state["last_vehicle_model"]

    # ✅ ต้องมีบรรทัดนี้ก่อนใช้ col1 !!!
    col1, col2, col3 = st.columns(3)
            
    # ----- ประเภท -----
    v_types = ["รถยนต์นั่งส่วนบุคคล","รถจักรยานยนต์","รถกระบะ","รถตู้","อื่นๆ"]
    def_type = st.session_state.form.get('vehicle_type','รถยนต์นั่งส่วนบุคคล')

    vehicle_type_sel = col1.selectbox(
        "ประเภทรถ",
        v_types,
        index=v_types.index(def_type) if def_type in v_types else 0,
        key="sel_vtype_t3"
    )

    vehicle_type_custom = ""
    if vehicle_type_sel == "อื่นๆ":
        vehicle_type_custom = col1.text_input(
            "ระบุประเภทรถ",
            st.session_state.form.get("vehicle_type_custom",""),
            key="input_vtype_custom_t3"
        )

    vehicle_type = vehicle_type_custom if vehicle_type_sel == "อื่นๆ" else vehicle_type_sel

    # ----- ยี่ห้อ -----
    v_brands = ["Toyota","Honda","Isuzu","Nissan","Mazda","Yamaha","อื่นๆ"]
    def_brand = st.session_state.form.get('vehicle_brand','Toyota')

    vehicle_brand_sel = col2.selectbox(
        "ยี่ห้อรถ",
        v_brands,
        index=v_brands.index(def_brand) if def_brand in v_brands else 0,
        key="sel_vbrand_t3"
    )

    vehicle_brand_custom = ""
    if vehicle_brand_sel == "อื่นๆ":
        vehicle_brand_custom = col2.text_input(
            "ระบุยี่ห้อรถ",
            st.session_state.form.get("vehicle_brand_custom",""),
            key="input_vbrand_custom_t3"
        )

    vehicle_brand = vehicle_brand_custom if vehicle_brand_sel == "อื่นๆ" else vehicle_brand_sel

        # ================= MODEL (SMART) =================
    st.write("---")
    st.caption("🚘 รุ่นรถ")

    # 🔥 เอารุ่นตามยี่ห้อ
    brand_for_model = vehicle_brand

    model_list = VEHICLE_MODELS.get(brand_for_model, [])

    # 👉 ถ้าไม่มีในระบบ ให้ fallback
    if not model_list:
        model_list = ["ไม่ทราบรุ่น", "อื่นๆ"]
    else:
        model_list = model_list + ["อื่นๆ"]

    # 👉 default ค่าเดิม
    def_model = st.session_state.form.get('vehicle_model', model_list[0])

    vehicle_model_sel = st.selectbox(
        "รุ่นรถ",
        model_list,
        index=model_list.index(def_model) if def_model in model_list else 0,
        key="sel_vmodel_t3"
    )

    vehicle_model_custom = ""

    if vehicle_model_sel == "อื่นๆ":
        vehicle_model_custom = st.text_input(
            "ระบุรุ่นรถ",
            st.session_state.form.get("vehicle_model_custom", ""),
            key="input_vmodel_custom_t3"
        )

    vehicle_model = vehicle_model_custom if vehicle_model_sel == "อื่นๆ" else vehicle_model_sel

    # ----- สี -----
    v_colors = ["ขาว","ดำ","เทา","แดง","น้ำเงิน","อื่นๆ"]
    def_color = st.session_state.form.get('vehicle_color','ขาว')

    vehicle_color_sel = col3.selectbox(
        "สีรถ",
        v_colors,
        index=v_colors.index(def_color) if def_color in v_colors else 0,
        key="sel_vcolor_t3"
    )

    vehicle_color_custom = ""
    if vehicle_color_sel == "อื่นๆ":
        vehicle_color_custom = col3.text_input(
            "ระบุสีรถ",
            st.session_state.form.get("vehicle_color_custom",""),
            key="input_vcolor_custom_t3"
        )

    vehicle_color = vehicle_color_custom if vehicle_color_sel == "อื่นๆ" else vehicle_color_sel

    # ================= ทะเบียนรถ =================
    st.write("---")
    st.caption("🔢 หมายเลขทะเบียน")

    plate_option = st.selectbox(
        "สถานะแผ่นป้ายทะเบียน",
        ["กรุณาเลือก", "มีแผ่นป้ายทะเบียน", "ไม่ติดแผ่นป้ายทะเบียน"],
        index=0,
        key="sel_plate_option_t3"
    )

    plate_letter = ""
    plate_number = ""
    plate_province = ""
    vehicle_plate = ""

    if plate_option == "มีแผ่นป้ายทะเบียน":

        c1, c2, c3 = st.columns([1,2,2])

        plate_letter = c1.text_input("ตัวอักษร", "", key="input_pletter_t3")
        plate_number = c2.text_input("ตัวเลข", "", key="input_pnumber_t3")

        def_plate = st.session_state.form.get('plate_province', provinces[0])

        plate_province = c3.selectbox(
            "จังหวัด",
            provinces,
            index=provinces.index(def_plate) if def_plate in provinces else 0,
            key="sel_pprov_t3"
        )

        vehicle_plate = f"{plate_letter} {plate_number} {plate_province}"
        st.success(f"ทะเบียน: {vehicle_plate}")

    elif plate_option == "ไม่ติดแผ่นป้ายทะเบียน":
        vehicle_plate = "ไม่ติดแผ่นป้ายทะเบียน"
        st.warning("🚫 ไม่มีแผ่นป้ายทะเบียน")

    else:
        vehicle_plate = ""
        st.info("ℹ️ กรุณาเลือกสถานะทะเบียน")

    # ================= LICENSE =================
    st.write("---")
    st.caption("🪪 ใบขับขี่")

    license_options = ["มี", "ไม่มี", "มี (แต่สิ้นอายุ)"]

    def_license = st.session_state.form.get("license_status", "มี")

    license_status = st.selectbox(
        "ใบขับขี่",
        license_options,
        index=license_options.index(def_license) if def_license in license_options else 0,
        key="sel_license_t3"
    )

    # ================= CASE =================
    st.write("---")
    st.caption("⚖️ ข้อหา")

    charge_list = [
        "เป็นผู้ขับขี่ (รถยนต์) ในขณะเมาสุรา",
        "เป็นผู้ขับขี่ (จักรยานยนต์) ในขณะเมาสุรา",
        "อื่นๆ"
    ]

    charge_sel = st.selectbox("ข้อหา", charge_list, key="sel_charge_t3")

    charge_custom = ""
    if charge_sel == "อื่นๆ":
        charge_custom = st.text_input("กรอกข้อหา", "", key="input_charge_custom_t3")

    final_charge = charge_custom if charge_sel == "อื่นๆ" else charge_sel

    # ================= DEVICE =================
    device_brand_list = ["ไลออน อัลโคลมิเตอร์", "อื่นๆ"]

    device_brand_sel = st.selectbox("ยี่ห้อเครื่องวัด", device_brand_list, key="sel_device_brand_t3")

    device_brand_custom = ""
    if device_brand_sel == "อื่นๆ":
        device_brand_custom = st.text_input("กรอกยี่ห้อ", "", key="input_device_brand_custom_t3")

    final_device_brand = device_brand_custom if device_brand_sel == "อื่นๆ" else device_brand_sel

    device_serial_list = ["T29542", "อื่นๆ"]

    device_serial_sel = st.selectbox("หมายเลขเครื่อง", device_serial_list, key="sel_device_serial_t3")

    device_serial_custom = ""
    if device_serial_sel == "อื่นๆ":
        device_serial_custom = st.text_input("กรอกหมายเลขเครื่อง", "", key="input_device_serial_custom_t3")

    final_device_serial = device_serial_custom if device_serial_sel == "อื่นๆ" else device_serial_sel

    alcohol = st.number_input("แอลกอฮอล์ (มก.%)", value=0, key="num_alc_t3")

    # ================= UPDATE =================
    st.session_state.form.update({
        "name": name,
        "sub_id": sub_id,
        "age": age,
        "house": house,
        "province": province,
        "amphur": amphur,
        "tumbon": tumbon,
        "vehicle_plate": vehicle_plate,
        "plate_letter": plate_letter,
        "plate_number": plate_number,
        "plate_province": plate_province,

        "vehicle_type": vehicle_type,
        "vehicle_brand": vehicle_brand,
        "vehicle_model": vehicle_model,
        "vehicle_color": vehicle_color,

        "plate_option": plate_option,

        "license_status": license_status,

        "charge": final_charge,

        "device_brand": final_device_brand,
        "device_serial": final_device_serial,

        "alcohol": alcohol
    })
    st.session_state["last_vehicle_brand"] = vehicle_brand
    st.session_state["last_vehicle_model"] = vehicle_model
    c1, c2 = st.columns(2)

    if c1.button("⬅️ ย้อนกลับ", key="btn_back_t3"):
        st.session_state.tab = 1
        st.rerun()

    if c2.button("ถัดไป ➡️", key="btn_next_t3"):
        st.session_state.tab = 3
        st.rerun()

# ================= TAB 4 =================
if st.session_state.tab == 3:
    st.write("")
    f = st.session_state.form
    
    # ================= FUNCTION: Thai Number =================
    def to_arabic_number(text):
        if text is None:
            return ""

        thai_to_arabic = {
            "๐":"0","๑":"1","๒":"2","๓":"3","๔":"4",
            "๕":"5","๖":"6","๗":"7","๘":"8","๙":"9"
        }

        text = str(text)

        return "".join(thai_to_arabic.get(ch, ch) for ch in text)
    
    # ================= จุดเกิดเหตุ =================
    st.caption("📌 จุดเกิดเหตุ")

    spot_options = [
        "บริเวณสี่แยกตู้ยามตระการร่วมใจ (เขตเทศบาลตระการพืชผล) หมู่ 4",
        "บริเวณสำนักงานเขตพื้นที่การศึกษาประถมศึกษาอุบลราชธานีเขต2 (เขตเทศบาลตำบลตระการพืชผล) หมู่4",
        "บริเวณหน้าสำนักงานที่ดินอำเภอตระการพืชผล (เขตเทศบาลตระการพืชผล) หมู่ 2",
        "บริเวณสี่แยกธนาคารกรุงเทพตระการพืชผล(เขตเทศบาลตำบลตระการพืชผล) หมู่ 2",
        "อื่นๆ"
    ]

    current_spot = f.get("spot_location", spot_options[0])

    selected_spot = st.selectbox(
        "เลือกจุดเกิดเหตุ",
        spot_options,
        index=spot_options.index(current_spot) if current_spot in spot_options else 0,
        key="sel_spot_location"
    )

    if selected_spot == "อื่นๆ":
        spot_custom = st.text_input(
            "ระบุจุดเกิดเหตุ",
            value=f.get("spot_location_custom", "")
        )
        selected_spot = spot_custom

    st.session_state.form["spot_location"] = selected_spot

    st.caption("📍 รายละเอียดสถานที่เกิดเหตุ")

    # ================= จุดเกิดเหตุ =================
    loc_options = [
        "ตั้งจุดตรวจบริเวณสี่แยกตู้ยามตระการร่วมใจ หมู่ 4",
        "ตั้งจุดตรวจบริเวณสำนักงานเขตพื้นที่การศึกษาประถมศึกษาอุบลราชธานีเขต2 หมู่4",
        "ตั้งจุดตรวจบริเวณหน้าสำนักงานที่ดินอำเภอตระการพืชผล หมู่ 2",
        "อำนวยความสะดวกการจราจรบริเวณสี่แยกตู้ยามตระการร่วมใจ หมู่ 4",
        "อำนวยความสะดวกการจราจรบริเวณสี่แยกธนาคารกรุงเทพตระการพืชผลหมู่ 2",
        "อื่นๆ"
    ]

    current_loc_detail = f.get('record_loc_detail', loc_options[0])
    default_loc_idx = loc_options.index(current_loc_detail) if current_loc_detail in loc_options else 0

    selected_loc = st.selectbox(
        "เลือกสถานที่เกิดเหตุ",
        loc_options,
        index=default_loc_idx
    )

    if selected_loc == "อื่นๆ":
        final_loc_detail = st.text_input(
            "ระบุสถานที่เกิดเหตุอื่นๆ",
            f.get('record_loc_detail_custom', '')
        )
    else:
        final_loc_detail = selected_loc

    # ================= จังหวัด/อำเภอ/ตำบล =================
    df_geo = load_geo()

    col_p, col_a, col_t = st.columns(3)

    provinces = sorted(df_geo['ProvinceThai'].dropna().unique())

    incident_province = col_p.selectbox(
        "จังหวัดที่เกิดเหตุ",
        provinces,
        index=provinces.index(f.get('incident_province')) if f.get('incident_province') in provinces else 0,
        key="incident_province"
    )

    districts = sorted(
        df_geo[df_geo['ProvinceThai'] == incident_province]['DistrictThai'].dropna().unique()
    )

    incident_amphur = col_a.selectbox(
        "อำเภอที่เกิดเหตุ",
        districts if districts else ["-"],
        index=districts.index(f.get('incident_amphur')) if f.get('incident_amphur') in districts else 0,
        key="incident_amphur"
    )

    sub_districts = sorted(
        df_geo[
            (df_geo['ProvinceThai'] == incident_province) &
            (df_geo['DistrictThai'] == incident_amphur)
        ]['TambonThai'].dropna().unique()
    )

    incident_tumbon = col_t.selectbox(
        "ตำบลที่เกิดเหตุ",
        sub_districts if sub_districts else ["-"],
        index=sub_districts.index(f.get('incident_tumbon')) if f.get('incident_tumbon') in sub_districts else 0,
        key="incident_tumbon"
    )

    st.write("---")

    # ================= จุดเชิญตัว =================
    invite_options = ["ตู้ยามตระการร่วมใจ", "สภ.ตระการพืชผล", "อื่นๆ"]
    current_invite = f.get('invite_loc', invite_options[0])
    default_invite_idx = invite_options.index(current_invite) if current_invite in invite_options else 0

    selected_invite = st.selectbox(
        "📍 จุดเชิญตัวผู้ต้องหา",
        invite_options,
        index=default_invite_idx
    )

    if selected_invite == "อื่นๆ":
        invite_loc = st.text_input("ระบุจุดเชิญตัวอื่นๆ", f.get('invite_loc_custom', ''))
    else:
        invite_loc = selected_invite

    confession_option = st.radio(
        "คำให้การผู้ต้องหา:",
        ["รับสารภาพตลอดข้อกล่าวหา", "ให้การปฏิเสธทุกข้อกล่าวหา"],
        index=0 if f.get('confession_status') == "รับสารภาพ" else 1,
        horizontal=True
    )

    col_d_manual, col_t_manual = st.columns(2)

    arrest_date_val = col_d_manual.date_input(
        "เมื่อวันที่ (แจ้งสิทธิ์)",
        value=f.get("arrest_date_manual", datetime.date.today())
    )

    arrest_time_val = col_t_manual.time_input(
        "เวลาประมาณ (แจ้งสิทธิ์)",
        value=f.get("arrest_time_manual", datetime.datetime.now().time())
    )

    st.caption("👮 ผู้เกี่ยวข้องในคดี")
    col_acc, col_wit, col_reader, col_proc = st.columns([2, 2, 2, 1.5])

    accuser_name = col_acc.selectbox("ผู้กล่าวหา", OFFICERS, index=0)
    witness_name = col_wit.selectbox("พยาน", OFFICERS, index=0)
    selected_officers = st.session_state.form.get("selected", [])

    # 👉 กันกรณีว่าง / หรือถูกเลือกหมด
    available_readers = [
        o for o in OFFICERS
        if o not in selected_officers
    ]

    # 🔥 fallback สำคัญ (กัน No options)
    if not available_readers:
        available_readers = OFFICERS.copy()
        st.warning("⚠️ ไม่มีรายชื่อว่าง ระบบจะให้เลือกทั้งหมดแทน")

    reader_name = col_reader.selectbox(
        "พยาน/ผู้บันทึกอ่าน",
        available_readers,
        key="reader_select_t4"
    )

    prosecutor_options = ["082-921-211", "อื่นๆ"]
    current_phone = f.get("prosecutor_phone", "082-921-211")
    phone_index = prosecutor_options.index(current_phone) if current_phone in prosecutor_options else 0

    selected_phone = col_proc.selectbox("📞 เบอร์อัยการ", prosecutor_options, index=phone_index)

    if selected_phone == "อื่นๆ":
        prosecutor_phone = st.text_input("กรอกเบอร์อัยการ", f.get("prosecutor_phone_custom", ""))
    else:
        prosecutor_phone = selected_phone

    st.write("---")

    final_detail = st.text_area(
        "ตรวจสอบ/แก้ไข พฤติการณ์จับกุม",
        value=f.get("detail", ""),
        height=250
    )

    # ================= UPDATE SESSION =================
    st.session_state.form.update({
        "accuser_name": accuser_name,
        "witness_name": witness_name,
        "reader_name": reader_name,
        "prosecutor_phone": prosecutor_phone,
        "detail": final_detail,
        "confession_status": confession_option,
        "arrest_date_manual": arrest_date_val,
        "arrest_time_manual": arrest_time_val,
        "record_loc_detail": final_loc_detail,
        "invite_loc": invite_loc,
        "incident_province": incident_province,
        "incident_amphur": incident_amphur,
        "incident_tumbon": incident_tumbon,
        "spot_location_custom": f.get("spot_location_custom", "")
    })

    # ================= RESET =================
    if st.session_state.get("confirm_reset", False):
        st.warning("⚠️ ต้องการล้างข้อมูลทั้งหมดใช่หรือไม่?")

        c1, c2 = st.columns(2)

        if c1.button("✅ ยืนยัน"):
            st.session_state.form = {}
            st.session_state.tab = 0
            st.session_state.confirm_reset = False
            st.rerun()

        if c2.button("❌ ยกเลิก"):
            st.session_state.confirm_reset = False
            st.rerun()

# ================= EXPORT =================
if st.session_state.tab == 3:
    st.write("---")

    left, right = st.columns(2)

    with left:
        if st.button("🔄 เริ่มต้นบันทึกใหม่", key="reset_btn", use_container_width=True):
            st.session_state.confirm_reset = True

    with right:
        if st.button("💾 สร้างเอกสารWord", key="export_btn", use_container_width=True):
            st.session_state.export_now = True

    if st.session_state.get("export_now", False):
        f = st.session_state.form
        required_fields = {
            "ชื่อผู้ต้องหา": f.get("name", "").strip(),
            "ข้อหา": f.get("charge", "").strip(),
            "ทะเบียนรถ": f.get("vehicle_plate", "").strip(),
        }
        missing = [k for k, v in required_fields.items() if not v]
        if missing:
            st.error(f"❌ กรุณากรอกข้อมูลให้ครบ: {', '.join(missing)}")
            st.session_state.export_now = False
        else:
            with st.spinner("⏳ กำลังสร้างเอกสาร..."):
                try:
                    selected = f.get("selected", [])
                    reader = f.get("reader_name", "")
                    selected_clean = [x for x in selected if x != reader]
                    sign_data = {}

                    for i in range(12):
                        if i < len(selected_clean):
                            sign_data[f"{{{{s{i+1}}}}}"] = "(ลงชื่อ) ................................................................."
                            sign_data[f"{{{{n{i+1}}}}}"] = f"({selected_clean[i]})"
                        else:
                            sign_data[f"{{{{s{i+1}}}}}"] = ""
                            sign_data[f"{{{{n{i+1}}}}}"] = ""

                    if reader:
                        sign_data["{{s13}}"] = "(ลงชื่อ) ................................................................."
                        sign_data["{{n13}}"] = f"({reader})"
                    else:
                        sign_data["{{s13}}"] = ""

                    vehicle_plate = f"{f.get('plate_letter','')}{f.get('plate_number','')} {f.get('plate_province','')}".strip()

                    incident_loc = " ".join(filter(None, [
                        f.get("spot_location", ""),
                        f"ต.{format_tambon(f.get('incident_tumbon',''))}" if f.get("incident_tumbon") else "",
                        f"อ.{format_amphur(f.get('incident_amphur',''))}" if f.get("incident_amphur") else "",
                        f"จ.{f.get('incident_province','')}" if f.get("incident_province") else ""
                    ])).strip()

                    def to_arabic_number(text):
                        if text is None:
                            return ""
                        thai_to_arabic = {
                            "๐":"0","๑":"1","๒":"2","๓":"3","๔":"4",
                            "๕":"5","๖":"6","๗":"7","๘":"8","๙":"9"
                        }
                        return "".join(thai_to_arabic.get(ch, ch) for ch in str(text))

                    data = {
                        "{{record_loc}}": f.get("record_loc", ""),
                        "{{location}}": f.get("record_loc_detail", ""),
                        "{{date}}": date_th(f.get("incident_date")),
                        "{{time}}": safe_time(f.get("arrest_time")),
                        "{{officer_name}}": f.get("accuser_name", ""),
                        "{{incident_date}}": date_th(f.get("incident_date")),
                        "{{arrest_time}}": safe_time(f.get("arrest_time")),
                        "{{report_date}}": date_th(f.get("report_date")),
                        "{{report_time}}": safe_time(f.get("report_time")),
                        "{{commander_1}}": f.get("commander_1", ""),
                        "{{commander_2}}": f.get("commander_2", ""),
                        "{{commander_3}}": f.get("commander_3", ""),
                        "{{officers_list}}": ", ".join(f.get("selected", [])),
                        "{{sub_name}}": f.get("name", ""),
                        "{{reader_name}}": f.get("reader_name", ""),
                        "{{age}}": str(f.get("age", "")),
                        "{{alcohol}}": to_arabic_number(f.get("alcohol", "")),
                        "{{sub_id}}": format_thai_id(f.get("sub_id", "")),
                        "{{sub_full_addr}}": f"{f.get('house','')} ต.{format_tambon(f.get('tumbon',''))} อ.{format_amphur(f.get('amphur',''))} จ.{f.get('province','')}",
                        "{{charge}}": f.get("charge", ""),
                        "{{incident_loc}}": incident_loc,
                        "{{incident_loc_detail}}": f.get("record_loc_detail", ""),
                        "{{tumbon}}": clean_geo(f.get("incident_tumbon", "")),
                        "{{amphur}}": clean_geo(f.get("incident_amphur", "")),
                        "{{province}}": clean_geo(f.get("incident_province", "")),
                        "{{incident_full_location}}": " ".join(filter(None, [
                            f"ต.{clean_geo(f.get('incident_tumbon',''))}" if f.get("incident_tumbon") else "",
                            f"อ.{format_amphur(f.get('incident_amphur',''))}" if f.get("incident_amphur") else "",
                            f"จ.{clean_geo(f.get('incident_province',''))}" if f.get("incident_province") else ""
                        ])).strip(),
                        "{{vehicle_type}}": f.get("vehicle_type", ""),
                        "{{vehicle_brand}}": f.get("vehicle_brand", ""),
                        "{{vehicle_model}}": f.get("vehicle_model", "-"),
                        "{{vehicle_color}}": f.get("vehicle_color", ""),
                        "{{vehicle_plate}}": vehicle_plate,
                        "{{license_status}}": f.get("license_status", ""),
                        "{{spot_location}}": " ".join(filter(None, [
                            f.get("spot_location", ""),
                            f"ต.{format_tambon(f.get('incident_tumbon',''))}" if f.get("incident_tumbon") else "",
                            f"อ.{format_amphur(f.get('incident_amphur',''))}" if f.get("incident_amphur") else "",
                            f"จ.{f.get('incident_province','')}" if f.get("incident_province") else ""
                        ])).strip(),
                        "{{invite_loc}}": f.get("invite_loc", ""),
                        "{{device_brand}}": f.get("device_brand", ""),
                        "{{device_serial}}": f.get("device_serial", ""),
                        "{{confession}}": f.get("confession_status", ""),
                        "{{statement}}": " ".join([
                            str(f.get("confession_status", "")),
                            str(f.get("detail", ""))
                        ]).strip(),
                        "{{send_date}}": date_th(f.get("report_date")),
                        "{{prosecutor_phone}}": f.get("prosecutor_phone", ""),
                        "{{accuser_name}}": f.get("accuser_name", ""),
                        "{{witness_name}}": f.get("witness_name", ""),
                    }

                    data.update(sign_data)
                    for k, v in data.items():
                        data[k] = to_arabic_number(v)

                    if not os.path.exists("template.docx"):
                        st.error("❌ ไม่พบไฟล์ template.docx")
                        st.session_state.export_now = False
                        st.stop()

                    doc = Document("template.docx")
                    replace_text(doc, data)

                    for p in doc.paragraphs:
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.line_spacing = 1

                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    p.paragraph_format.space_before = Pt(0)
                                    p.paragraph_format.space_after = Pt(0)
                                    p.paragraph_format.line_spacing = 1

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    st.download_button(
                        "📥 ดาวน์โหลดWord",
                        buffer,
                        file_name=f"{f.get('name','arrest_report')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.session_state.export_now = False

                except Exception as e:
                    st.error(f"❌ {e}")
                    st.session_state.export_now = False
